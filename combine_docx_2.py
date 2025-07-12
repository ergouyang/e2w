from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os

def merge_word_documents(file_paths, output_path):
    # 创建目标文档
    target_doc = Document()
    
    # 移除目标文档的默认空白段落
    if len(target_doc.paragraphs) > 0:
        target_doc._body.clear_content()
    
    for i, file_path in enumerate(file_paths):
        # 打开源文档
        source_doc = Document(file_path)
        
        # 复制所有内容
        for element in source_doc.element.body:
            # 深度复制XML元素
            new_element = OxmlElement(element.tag)
            new_element[:] = [child for child in element]
            target_doc.element.body.append(new_element)
        
        # 在文档结尾添加分节符（除最后一个文档外）
        if i < len(file_paths) - 1:
            add_section_break(target_doc)
    
    # 保存合并后的文档
    target_doc.save(output_path)

def add_section_break(doc):
    """添加下一页分节符"""
    p = doc.add_paragraph()
    run = p.add_run()
    break_element = OxmlElement('w:br')
    break_element.set(qn('w:type'), 'page')
    run._r.append(break_element)
    
    # 添加分节符
    sect_pr = OxmlElement('w:pPr')
    p._p.append(sect_pr)
    sect_break = OxmlElement('w:sectPr')
    sect_pr.append(sect_break)

def batch_merge_word(folder_path, output_file):
    # 获取文件夹中所有docx文件
    file_list = [os.path.join(folder_path, f) for f in os.listdir(folder_path) 
                 if f.endswith('.docx')]
    
    if not file_list:
        print("未找到Word文档")
        return
    
    # 合并文档
    merge_word_documents(file_list, output_file)
    print(f"成功合并 {len(file_list)} 个文档到 {output_file}")

# 使用示例
if __name__ == "__main__":
    # 设置输入文件夹和输出路径
    input_folder = r"E:\e2w\大同-天津南-20250712\res2"  # 替换为你的文件夹路径
    output_file = r"E:\e2w\大同-天津南-20250712\re\merged_document.docx"
    
    batch_merge_word(input_folder, output_file)
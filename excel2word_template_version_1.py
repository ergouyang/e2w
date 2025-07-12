#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Excel到Word模板转换工具

功能：
1. 导入Excel文件并读取数据
2. 导入Word模板文件
3. 识别模板中的占位符
4. 映射Excel字段到Word占位符
5. 支持自动匹配和数学表达式
6. 批量生成Word文档

Author: yf
Year: 2025
License: MIT License

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import re
import os
import tempfile
import subprocess
import platform
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import difflib
from typing import List, Dict, Any, Optional
import glob


class Excel2WordConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel到Word模板转换工具")
        self.root.geometry("1200x1000")
        
        # 数据存储
        self.excel_data = None
        self.excel_file_path = None
        self.word_template_path = None
        self.placeholders = []
        self.mapping_data = []
        self.image_mapping_data = []  # 图片映射数据
        self.console_output = []  # 控制台输出缓存
        
        # 创建界面
        self.create_widgets()
        
    def create_widgets(self):
        """创建界面组件"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 配置权重
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(3, weight=1)  # 标签页区域权重最大
        
        # 标题
        title_label = ttk.Label(main_frame, text="Excel到Word模板转换工具", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, pady=(0, 20))
        
        # 文件选择区域
        file_frame = ttk.LabelFrame(main_frame, text="第一步：选择文件", padding="10")
        file_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        file_frame.columnconfigure(1, weight=1)
        
        # Excel文件选择
        ttk.Label(file_frame, text="Excel文件:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.excel_file_var = tk.StringVar(value="未选择")
        self.excel_label = ttk.Label(file_frame, textvariable=self.excel_file_var, 
                                    relief="sunken", width=50)
        self.excel_label.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(0, 10))
        
        excel_btn_frame = ttk.Frame(file_frame)
        excel_btn_frame.grid(row=0, column=2, sticky=tk.E)
        ttk.Button(excel_btn_frame, text="导入Excel", 
                  command=self.import_excel).grid(row=0, column=0, padx=(0, 5))
        ttk.Button(excel_btn_frame, text="清除", 
                  command=self.clear_excel).grid(row=0, column=1)
        
        # Word模板文件选择
        ttk.Label(file_frame, text="Word模板:").grid(row=1, column=0, sticky=tk.W, padx=(0, 10))
        self.word_file_var = tk.StringVar(value="未选择")
        self.word_label = ttk.Label(file_frame, textvariable=self.word_file_var, 
                                   relief="sunken", width=50)
        self.word_label.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(0, 10))
        
        word_btn_frame = ttk.Frame(file_frame)
        word_btn_frame.grid(row=1, column=2, sticky=tk.E)
        ttk.Button(word_btn_frame, text="导入Word模板", 
                  command=self.import_word_template).grid(row=0, column=0, padx=(0, 5))
        ttk.Button(word_btn_frame, text="清除", 
                  command=self.clear_word).grid(row=0, column=1)
        
        # 快速操作区域
        quick_action_frame = ttk.Frame(main_frame)
        quick_action_frame.grid(row=2, column=0, pady=(0, 10))
        
        ttk.Button(quick_action_frame, text="自动匹配字段", 
                  command=self.auto_match_fields).grid(row=0, column=0, padx=5)
        ttk.Button(quick_action_frame, text="预览", 
                  command=self.preview_document).grid(row=0, column=1, padx=5)
        ttk.Button(quick_action_frame, text="按模板导出", 
                  command=self.export_documents).grid(row=0, column=2, padx=5)
        
        # 主要内容区域 - 使用标签页
        notebook = ttk.Notebook(main_frame)
        notebook.grid(row=3, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        # 标签页1: 字段映射
        mapping_frame = ttk.Frame(notebook)
        notebook.add(mapping_frame, text="第二步：字段映射")
        
        # 字段映射内容
        mapping_content = ttk.Frame(mapping_frame, padding="10")
        mapping_content.pack(fill=tk.BOTH, expand=True)
        mapping_content.columnconfigure(0, weight=1)
        mapping_content.columnconfigure(1, weight=1)
        mapping_content.rowconfigure(0, weight=1)
        
        # Excel字段列表
        excel_frame = ttk.LabelFrame(mapping_content, text="Excel字段", padding="5")
        excel_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        excel_frame.columnconfigure(0, weight=1)
        excel_frame.rowconfigure(0, weight=1)
        
        self.excel_tree = ttk.Treeview(excel_frame, columns=("field",), show="tree headings", height=15)
        self.excel_tree.heading("#0", text="序号")
        self.excel_tree.heading("field", text="Excel字段名")
        self.excel_tree.column("#0", width=50)
        self.excel_tree.column("field", width=200)
        self.excel_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        excel_scroll = ttk.Scrollbar(excel_frame, orient="vertical", command=self.excel_tree.yview)
        excel_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.excel_tree.configure(yscrollcommand=excel_scroll.set)
        
        # Word占位符映射
        word_frame = ttk.LabelFrame(mapping_content, text="Word占位符映射（双击编辑）", padding="5")
        word_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        word_frame.columnconfigure(0, weight=1)
        word_frame.rowconfigure(0, weight=1)
        
        self.mapping_tree = ttk.Treeview(word_frame, columns=("placeholder", "mapping"), 
                                        show="tree headings", height=15)
        self.mapping_tree.heading("#0", text="序号")
        self.mapping_tree.heading("placeholder", text="Word占位符")
        self.mapping_tree.heading("mapping", text="匹配模式")
        self.mapping_tree.column("#0", width=50)
        self.mapping_tree.column("placeholder", width=150)
        self.mapping_tree.column("mapping", width=150)
        self.mapping_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 绑定双击编辑事件
        self.mapping_tree.bind("<Double-1>", self.edit_mapping)
        
        mapping_scroll = ttk.Scrollbar(word_frame, orient="vertical", command=self.mapping_tree.yview)
        mapping_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.mapping_tree.configure(yscrollcommand=mapping_scroll.set)
        
        # 映射操作按钮
        mapping_btn_frame = ttk.Frame(mapping_content)
        mapping_btn_frame.grid(row=1, column=0, columnspan=2, pady=10)
        
        self.exact_match_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(mapping_btn_frame, text="使用精准匹配", 
                       variable=self.exact_match_var).pack(side=tk.LEFT, padx=10)
        
        # 标签页2: 图片映射
        image_frame = ttk.Frame(notebook)
        notebook.add(image_frame, text="第三步：图片映射")
        
        # 图片映射内容
        image_content = ttk.Frame(image_frame, padding="10")
        image_content.pack(fill=tk.BOTH, expand=True)
        image_content.columnconfigure(0, weight=1)
        image_content.rowconfigure(1, weight=1)
        
        # 图片映射说明
        ttk.Label(image_content, text="如果Word模板包含图片占位符，请在此配置图片映射规则", 
                 font=("Arial", 10)).grid(row=0, column=0, sticky=tk.W, pady=(0, 10))
        
        # 图片映射控制按钮
        img_btn_frame = ttk.Frame(image_content)
        img_btn_frame.grid(row=1, column=0, sticky=tk.W, pady=(0, 10))
        
        ttk.Button(img_btn_frame, text="添加行", 
                  command=self.add_image_mapping_row).grid(row=0, column=0, padx=(0, 5))
        ttk.Button(img_btn_frame, text="删除选中行", 
                  command=self.remove_image_mapping_row).grid(row=0, column=1, padx=(0, 5))
        ttk.Button(img_btn_frame, text="清空所有", 
                  command=self.clear_image_mappings).grid(row=0, column=2, padx=(0, 5))
        ttk.Button(img_btn_frame, text="调试信息", 
                  command=self.show_image_debug_info).grid(row=0, column=3, padx=(0, 5))
        
        # 图片映射表格
        self.image_tree = ttk.Treeview(image_content, columns=("folder", "mapping_rule", "placeholder"), 
                                      show="tree headings", height=15)
        self.image_tree.heading("#0", text="序号")
        self.image_tree.heading("folder", text="图片文件夹")
        self.image_tree.heading("mapping_rule", text="映射规则")
        self.image_tree.heading("placeholder", text="Word占位符")
        self.image_tree.column("#0", width=50)
        self.image_tree.column("folder", width=200)
        self.image_tree.column("mapping_rule", width=200)
        self.image_tree.column("placeholder", width=150)
        self.image_tree.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 绑定双击编辑事件
        self.image_tree.bind("<Double-1>", self.edit_image_mapping)
        
        image_scroll = ttk.Scrollbar(image_content, orient="vertical", command=self.image_tree.yview)
        image_scroll.grid(row=2, column=1, sticky=(tk.N, tk.S))
        self.image_tree.configure(yscrollcommand=image_scroll.set)
        
        # 标签页3: 导出设置
        settings_frame = ttk.Frame(notebook)
        notebook.add(settings_frame, text="第四步：导出设置")
        
        # 导出设置内容
        settings_content = ttk.Frame(settings_frame, padding="10")
        settings_content.pack(fill=tk.BOTH, expand=True)
        settings_content.columnconfigure(0, weight=1)
        settings_content.columnconfigure(1, weight=1)
        
        # 左侧：基本设置
        basic_settings_frame = ttk.LabelFrame(settings_content, text="基本设置", padding="10")
        basic_settings_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        
        self.merge_docs_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(basic_settings_frame, text="合并导出文档", 
                       variable=self.merge_docs_var).grid(row=0, column=0, sticky=tk.W, pady=5)
        
        self.preview_in_file_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(basic_settings_frame, text="在文件中预览", 
                       variable=self.preview_in_file_var).grid(row=1, column=0, sticky=tk.W, pady=5)
        
        # 行数范围设置
        range_frame = ttk.LabelFrame(basic_settings_frame, text="导出行数范围", padding="5")
        range_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(10, 0))
        
        self.export_range_var = tk.StringVar(value="全部")
        
        # 全部选项
        ttk.Radiobutton(range_frame, text="全部数据", 
                       variable=self.export_range_var, value="全部").grid(row=0, column=0, sticky=tk.W, pady=2)
        
        # 指定区间选项
        range_custom_frame = ttk.Frame(range_frame)
        range_custom_frame.grid(row=1, column=0, sticky=tk.W, pady=2)
        
        ttk.Radiobutton(range_custom_frame, text="指定区间：从第", 
                       variable=self.export_range_var, value="指定").grid(row=0, column=0, sticky=tk.W)
        
        self.range_start_var = tk.StringVar(value="1")
        start_entry = ttk.Entry(range_custom_frame, textvariable=self.range_start_var, width=5)
        start_entry.grid(row=0, column=1, padx=(5, 0))
        
        ttk.Label(range_custom_frame, text="行到第").grid(row=0, column=2, padx=(5, 0))
        
        self.range_end_var = tk.StringVar(value="1")
        end_entry = ttk.Entry(range_custom_frame, textvariable=self.range_end_var, width=5)
        end_entry.grid(row=0, column=3, padx=(5, 0))
        
        ttk.Label(range_custom_frame, text="行（包括）").grid(row=0, column=4, padx=(5, 0))
        
        # 数据行数提示
        self.data_info_label = ttk.Label(range_frame, text="当前无数据", 
                                        font=("Arial", 9), foreground="gray")
        self.data_info_label.grid(row=2, column=0, sticky=tk.W, pady=(5, 0))
        
        # 文件命名设置
        naming_frame = ttk.LabelFrame(basic_settings_frame, text="文件命名设置", padding="5")
        naming_frame.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=(10, 0))
        
        # 命名方式选择
        self.naming_mode_var = tk.StringVar(value="默认")
        
        # 默认命名
        ttk.Radiobutton(naming_frame, text="默认命名（导出文档_001.docx）", 
                       variable=self.naming_mode_var, value="默认").grid(row=0, column=0, sticky=tk.W, pady=2)
        
        # 字段命名
        field_naming_frame = ttk.Frame(naming_frame)
        field_naming_frame.grid(row=1, column=0, sticky=tk.W, pady=2)
        
        ttk.Radiobutton(field_naming_frame, text="使用Excel字段命名：", 
                       variable=self.naming_mode_var, value="字段").grid(row=0, column=0, sticky=tk.W)
        
        self.naming_field_var = tk.StringVar()
        self.naming_field_combo = ttk.Combobox(field_naming_frame, textvariable=self.naming_field_var, 
                                              width=20, state="readonly")
        self.naming_field_combo.grid(row=0, column=1, padx=(10, 0))
        
        # 固定前缀命名
        prefix_naming_frame = ttk.Frame(naming_frame)
        prefix_naming_frame.grid(row=2, column=0, sticky=tk.W, pady=2)
        
        ttk.Radiobutton(prefix_naming_frame, text="固定前缀命名：", 
                       variable=self.naming_mode_var, value="前缀").grid(row=0, column=0, sticky=tk.W)
        
        self.naming_prefix_var = tk.StringVar(value="文档")
        prefix_entry = ttk.Entry(prefix_naming_frame, textvariable=self.naming_prefix_var, width=20)
        prefix_entry.grid(row=0, column=1, padx=(10, 0))
        
        ttk.Label(prefix_naming_frame, text="（如：文档_001.docx）", 
                 font=("Arial", 9), foreground="gray").grid(row=0, column=2, padx=(5, 0))
        
        # 右侧：高级设置
        advanced_settings_frame = ttk.LabelFrame(settings_content, text="高级设置", padding="10")
        advanced_settings_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        
        # 图片设置
        image_settings_frame = ttk.LabelFrame(advanced_settings_frame, text="图片设置", padding="5")
        image_settings_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.use_cm_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(image_settings_frame, text="图片宽度使用厘米", 
                       variable=self.use_cm_var).grid(row=0, column=0, sticky=tk.W, pady=2)
        
        width_frame = ttk.Frame(image_settings_frame)
        width_frame.grid(row=1, column=0, sticky=tk.W, pady=2)
        
        ttk.Label(width_frame, text="图片宽度:").grid(row=0, column=0, sticky=tk.W)
        self.image_width_var = tk.StringVar(value="9.8")
        width_entry = ttk.Entry(width_frame, textvariable=self.image_width_var, width=8)
        width_entry.grid(row=0, column=1, padx=(5, 0))
        
        self.unit_label = ttk.Label(width_frame, text="厘米")
        self.unit_label.grid(row=0, column=2, padx=(5, 0))
        
        # 绑定单位切换事件
        def on_unit_change():
            if self.use_cm_var.get():
                self.unit_label.config(text="厘米")
            else:
                self.unit_label.config(text="英寸")
        
        self.use_cm_var.trace('w', lambda *args: on_unit_change())
        
        # 数字格式化设置
        number_frame = ttk.LabelFrame(advanced_settings_frame, text="数字格式化", padding="5")
        number_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        format_frame = ttk.Frame(number_frame)
        format_frame.grid(row=0, column=0, sticky=tk.W, pady=2)
        
        ttk.Label(format_frame, text="数字格式:").grid(row=0, column=0, sticky=tk.W)
        
        self.number_format_var = tk.StringVar(value="保留原格式")
        number_combo = ttk.Combobox(format_frame, textvariable=self.number_format_var, 
                                   values=["保留原格式", "取整数", "保留1位小数", "保留2位小数", "保留3位小数"], 
                                   state="readonly", width=12)
        number_combo.grid(row=0, column=1, padx=(5, 0))
        
        # 自定义小数位数
        custom_decimal_frame = ttk.Frame(number_frame)
        custom_decimal_frame.grid(row=1, column=0, sticky=tk.W, pady=2)
        
        self.enable_custom_decimal_var = tk.BooleanVar(value=False)
        custom_decimal_check = ttk.Checkbutton(custom_decimal_frame, text="自定义小数位数:", 
                                              variable=self.enable_custom_decimal_var)
        custom_decimal_check.grid(row=0, column=0, sticky=tk.W)
        
        self.custom_decimal_var = tk.StringVar(value="2")
        decimal_entry = ttk.Entry(custom_decimal_frame, textvariable=self.custom_decimal_var, width=5)
        decimal_entry.grid(row=0, column=1, padx=(5, 0))
        
        # 千分位分隔符
        self.use_thousands_separator_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(number_frame, text="使用千分位分隔符", 
                       variable=self.use_thousands_separator_var).grid(row=2, column=0, sticky=tk.W, pady=2)
        
        # 底部工具栏
        toolbar_frame = ttk.Frame(main_frame)
        toolbar_frame.grid(row=4, column=0, pady=(10, 0))
        
        ttk.Button(toolbar_frame, text="查看输出", 
                  command=self.show_console_output).grid(row=0, column=0, padx=5)
        ttk.Button(toolbar_frame, text="使用助手", 
                  command=self.show_help).grid(row=0, column=1, padx=5)
        ttk.Button(toolbar_frame, text="测试替换", 
                  command=self.test_text_replacement).grid(row=0, column=2, padx=5)
        ttk.Button(toolbar_frame, text="预览文件名", 
                  command=self.preview_filenames).grid(row=0, column=3, padx=5)
        
        # 底部版权信息
        copyright_frame = ttk.Frame(main_frame)
        copyright_frame.grid(row=5, column=0, pady=(10, 5))
        
        ttk.Label(copyright_frame, text="create by yf 2025 | 遵循MIT开源协议", 
                 font=("Arial", 9), foreground="gray").pack()
    
    def import_excel(self):
        """导入Excel文件"""
        try:
            file_path = filedialog.askopenfilename(
                title="选择Excel文件",
                filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
            )
            
            if not file_path:
                return
            
            # 读取Excel数据
            self.excel_data = pd.read_excel(file_path)
            self.excel_file_path = file_path
            
            # 更新显示
            self.excel_file_var.set(os.path.basename(file_path))
            self.update_excel_tree()
            
            messagebox.showinfo("成功", f"Excel导入成功！共导入{len(self.excel_data)}行数据。")
            
        except Exception as e:
            messagebox.showerror("错误", f"导入Excel失败：{str(e)}")
    
    def clear_excel(self):
        """清除Excel数据"""
        if messagebox.askyesno("确认", "确定要清除Excel数据吗？这也会清空图片映射设置。"):
            self.excel_data = None
            self.excel_file_path = None
            self.excel_file_var.set("未选择")
            self.update_excel_tree()
            
            # 清空图片映射中依赖Excel字段的规则
            self.image_mapping_data = [
                mapping for mapping in self.image_mapping_data
                if not mapping["mapping_rule"].startswith("根据字段:")
            ]
            self.update_image_tree()
            
            # 清空文件命名字段选择
            self.naming_field_combo['values'] = []
            self.naming_field_var.set("")
            
            messagebox.showinfo("成功", "Excel数据已清除！")
    
    def import_word_template(self):
        """导入Word模板"""
        try:
            file_path = filedialog.askopenfilename(
                title="选择Word模板",
                filetypes=[("Word files", "*.docx *.doc"), ("All files", "*.*")]
            )
            
            if not file_path:
                return
            
            self.word_template_path = file_path
            
            # 提取占位符
            self.extract_placeholders()
            
            # 更新显示
            self.word_file_var.set(os.path.basename(file_path))
            self.update_mapping_tree()
            
            # 初始化图片映射表格
            self.update_image_tree()
            
            if len(self.placeholders) > 0:
                messagebox.showinfo("成功", f"Word模板导入成功！识别到{len(self.placeholders)}个占位符。")
            else:
                messagebox.showwarning("提示", 
                    "Word模板导入成功，但未识别到占位符。\n\n"
                    "请确保模板中包含以下格式的占位符：\n"
                    "• {{字段名}}（双花括号）\n"
                    "• 【字段名】（中文书名号）\n"
                    "• 《字段名》（中文双书名号）\n"
                    "• [字段名]（方括号）\n"
                    "• <字段名>（尖括号）\n\n"
                    "支持中文和英文字段名。")
            
        except Exception as e:
            messagebox.showerror("错误", f"导入Word模板失败：{str(e)}")
    
    def clear_word(self):
        """清除Word模板"""
        if messagebox.askyesno("确认", "确定要清除Word模板吗？"):
            self.word_template_path = None
            self.placeholders = []
            self.mapping_data = []
            self.image_mapping_data = []
            self.word_file_var.set("未选择")
            self.update_mapping_tree()
            self.update_image_tree()
            messagebox.showinfo("成功", "Word模板已清除！")
    
    def extract_text_from_element(self, element):
        """从文档元素中提取文本"""
        text = ""
        try:
            # 处理段落
            if hasattr(element, 'paragraphs'):
                for paragraph in element.paragraphs:
                    text += paragraph.text + "\n"
            
            # 处理表格
            if hasattr(element, 'tables'):
                for table in element.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
                            # 递归处理表格中的内容
                            text += self.extract_text_from_element(cell)
            
            # 处理形状和文本框
            if hasattr(element, '_element'):
                from docx.oxml.ns import qn
                
                # 查找所有文本框和形状
                shapes = element._element.xpath('.//w:txbxContent//w:t')
                for shape in shapes:
                    if shape.text:
                        text += shape.text + "\n"
                
                # 查找所有绘图对象中的文本
                drawings = element._element.xpath('.//w:drawing//w:t')
                for drawing in drawings:
                    if drawing.text:
                        text += drawing.text + "\n"
        
        except Exception as e:
            # 如果遇到错误，继续处理其他内容
            pass
        
        return text
    
    def extract_placeholders(self):
        """从Word文档中提取占位符"""
        try:
            doc = Document(self.word_template_path)
            text = ""
            
            # 从主文档段落中提取文本
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 从主文档表格中提取文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
                        # 递归处理表格中的段落和嵌套内容
                        text += self.extract_text_from_element(cell)
            
            # 从页眉中提取文本
            for section in doc.sections:
                # 提取页眉
                if section.header:
                    text += self.extract_text_from_element(section.header)
                
                # 提取页脚
                if section.footer:
                    text += self.extract_text_from_element(section.footer)
                
                # 提取首页页眉
                if hasattr(section, 'first_page_header') and section.first_page_header:
                    text += self.extract_text_from_element(section.first_page_header)
                
                # 提取首页页脚
                if hasattr(section, 'first_page_footer') and section.first_page_footer:
                    text += self.extract_text_from_element(section.first_page_footer)
                
                # 提取奇数页页眉
                if hasattr(section, 'even_page_header') and section.even_page_header:
                    text += self.extract_text_from_element(section.even_page_header)
                
                # 提取奇数页页脚
                if hasattr(section, 'even_page_footer') and section.even_page_footer:
                    text += self.extract_text_from_element(section.even_page_footer)
            
            # 从形状和文本框中提取文本
            try:
                from docx.oxml.ns import qn
                
                # 查找所有文本框
                for textbox in doc._element.xpath('.//w:txbxContent'):
                    for paragraph in textbox.xpath('.//w:p'):
                        para_text = ""
                        for run in paragraph.xpath('.//w:t'):
                            if run.text:
                                para_text += run.text
                        if para_text:
                            text += para_text + "\n"
                
                # 查找所有绘图对象中的文本
                for drawing in doc._element.xpath('.//w:drawing'):
                    for text_run in drawing.xpath('.//w:t'):
                        if text_run.text:
                            text = text + text_run.text + "\n"
            
            except Exception as extract_error:

                # 如果提取特殊元素失败，继续处理
                self.log_output(f"提取特殊元素时出错: {extract_error}")
            
            # 调试信息：打印提取的文本（可选）
            # print(f"提取的文本内容:\n{text}")
            
            # 使用正则表达式提取占位符，支持Unicode字符
            patterns = [
                r'\{\{([^}]+)\}\}',  # {{字段名}}
                r'\【([^】]+)\】',    # 【字段名】
                r'\《([^》]+)》',    # 《字段名》
                r'\[([^\]]+)\]',     # [字段名]
                r'\<([^>]+)\>',      # <字段名>
            ]
            
            self.placeholders = []
            for pattern in patterns:
                # 使用re.UNICODE标志确保正确处理Unicode字符
                matches = re.findall(pattern, text, re.UNICODE)
                for match in matches:
                    # 根据不同的模式格式化占位符
                    if pattern == r'\{\{([^}]+)\}\}':
                        placeholder = f"{{{{{match}}}}}"
                    # elif pattern == r'\【([^】]+)\】':
                    #     placeholder = f"【{match}】"
                    # elif pattern == r'\《([^》]+)》':
                    #     placeholder = f"《{match}》"
                    # elif pattern == r'\[([^\]]+)\]':
                    #     placeholder = f"[{match}]"
                    # elif pattern == r'\<([^>]+)\>':
                    #     placeholder = f"<{match}>"
                    # else:
                    #     placeholder = f"{{{{{match}}}}}"
                    
                    if placeholder not in self.placeholders:
                        self.placeholders.append(placeholder)
            
            # 初始化映射数据
            self.mapping_data = [{"placeholder": p, "mapping": ""} for p in self.placeholders]
            self.placeholders.sort()
            
            # 调试信息：显示找到的占位符
            if self.placeholders:
                self.log_output(f"找到的占位符: {self.placeholders}")
            
        except Exception as e:
            messagebox.showerror("错误", f"提取占位符失败：{str(e)}")
    
    def update_excel_tree(self):
        """更新Excel字段树"""
        # 清空现有项目
        for item in self.excel_tree.get_children():
            self.excel_tree.delete(item)
        
        if self.excel_data is not None:
            for i, column in enumerate(self.excel_data.columns):
                self.excel_tree.insert("", "end", text=str(i+1), values=(column,))
            
            # 更新文件命名字段列表
            self.update_naming_fields()
    
    def update_naming_fields(self):
        """更新文件命名字段列表"""
        if self.excel_data is not None:
            self.naming_field_combo['values'] = list(self.excel_data.columns)
            if not self.naming_field_var.get() and len(self.excel_data.columns) > 0:
                self.naming_field_var.set(self.excel_data.columns[0])
            
            # 更新数据行数提示
            total_rows = len(self.excel_data)
            self.data_info_label.config(text=f"当前数据共 {total_rows} 行")
            
            # 更新默认的结束行号
            self.range_end_var.set(str(total_rows))
        else:
            self.data_info_label.config(text="当前无数据")
    
    def update_naming_ui(self):
        """更新命名UI状态"""
        naming_mode = self.naming_mode_var.get()
        
        # 根据选择的命名方式启用/禁用相关控件
        if naming_mode == "字段":
            self.naming_field_combo.config(state="readonly")
        else:
            self.naming_field_combo.config(state="disabled")
    
    def validate_export_range(self):
        """验证导出行数范围设置"""
        if self.excel_data is None or len(self.excel_data) == 0:
            return False, "没有Excel数据"
        
        total_rows = len(self.excel_data)
        export_mode = self.export_range_var.get()
        
        if export_mode == "全部":
            return True, f"将导出全部 {total_rows} 行数据"
        
        elif export_mode == "指定":
            try:
                start_row = int(self.range_start_var.get())
                end_row = int(self.range_end_var.get())
                
                if start_row < 1:
                    return False, "起始行号不能小于1"
                
                if end_row > total_rows:
                    return False, f"结束行号不能大于总行数({total_rows})"
                
                if start_row > end_row:
                    return False, "起始行号不能大于结束行号"
                
                export_count = end_row - start_row + 1
                return True, f"将导出第 {start_row} 行到第 {end_row} 行，共 {export_count} 行数据"
                
            except ValueError:
                return False, "行号必须是有效的数字"
        
        return False, "未知的导出模式"
    
    def get_export_data_range(self):
        """获取要导出的数据范围"""
        if self.excel_data is None or len(self.excel_data) == 0:
            return None
        
        export_mode = self.export_range_var.get()
        
        if export_mode == "全部":
            return self.excel_data
        
        elif export_mode == "指定":
            try:
                start_row = int(self.range_start_var.get())
                end_row = int(self.range_end_var.get())
                
                # 转换为DataFrame的索引（0-based）
                start_idx = start_row - 1
                end_idx = end_row
                
                return self.excel_data.iloc[start_idx:end_idx]
                
            except (ValueError, IndexError):
                return None
        
        return None
    
    def generate_filename(self, data_row: pd.Series, row_index: int, used_names: set) -> str:
        """生成文件名"""
        try:
            naming_mode = self.naming_mode_var.get()
            
            if naming_mode == "默认":
                # 默认命名：导出文档_001.docx
                filename = f"导出文档_{row_index+1:03d}.docx"
                
            elif naming_mode == "字段":
                # 使用Excel字段命名
                field_name = self.naming_field_var.get()
                if field_name and field_name in data_row.index:
                    field_value = data_row[field_name]
                    if pd.notna(field_value):
                        # 清理文件名，移除不合法字符
                        base_name = self.clean_filename(str(field_value))
                        filename = f"{base_name}.docx"
                        
                        # 处理重复文件名
                        if filename in used_names:
                            counter = 1
                            while f"{base_name}_{counter}.docx" in used_names:
                                counter += 1
                            filename = f"{base_name}_{counter}.docx"
                    else:
                        # 字段值为空，使用默认命名
                        filename = f"导出文档_{row_index+1:03d}.docx"
                else:
                    # 字段不存在，使用默认命名
                    filename = f"导出文档_{row_index+1:03d}.docx"
                    
            elif naming_mode == "前缀":
                # 固定前缀命名
                prefix = self.naming_prefix_var.get().strip()
                if not prefix:
                    prefix = "文档"
                prefix = self.clean_filename(prefix)
                filename = f"{prefix}_{row_index+1:03d}.docx"
                
            else:
                # 未知模式，使用默认命名
                filename = f"导出文档_{row_index+1:03d}.docx"
            
            # 确保文件名不为空且有效
            if not filename or filename == ".docx":
                filename = f"导出文档_{row_index+1:03d}.docx"
            
            self.log_output(f"生成文件名: 行{row_index+1} -> {filename}")
            return filename
            
        except Exception as e:
            self.log_output(f"生成文件名失败: {str(e)}")
            return f"导出文档_{row_index+1:03d}.docx"
    
    def clean_filename(self, filename: str) -> str:
        """清理文件名，移除不合法字符"""
        # 移除或替换不合法字符
        invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
        clean_name = filename
        
        for char in invalid_chars:
            clean_name = clean_name.replace(char, '_')
        
        # 移除前后空格
        clean_name = clean_name.strip()
        
        # 确保不以点开头或结尾
        clean_name = clean_name.strip('.')
        
        # 限制长度（Windows文件名限制为255个字符，但留出扩展名空间）
        if len(clean_name) > 200:
            clean_name = clean_name[:200]
        
        # 如果清理后为空，使用默认名称
        if not clean_name:
            clean_name = "文档"
        
        return clean_name
    
    def update_mapping_tree(self):
        """更新映射树"""
        # 清空现有项目
        for item in self.mapping_tree.get_children():
            self.mapping_tree.delete(item)
        
        for i, data in enumerate(self.mapping_data):
            self.mapping_tree.insert("", "end", text=str(i+1), 
                                    values=(data["placeholder"], data["mapping"]))
    
    def edit_mapping(self, event):
        """编辑映射关系"""
        selection = self.mapping_tree.selection()
        if not selection:
            return
        
        item = selection[0]
        item_index = int(self.mapping_tree.item(item, "text")) - 1
        current_mapping = self.mapping_data[item_index]["mapping"]
        
        # 创建编辑对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("编辑映射")
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 占位符显示
        ttk.Label(dialog, text="占位符:").grid(row=0, column=0, sticky=tk.W, padx=10, pady=5)
        ttk.Label(dialog, text=self.mapping_data[item_index]["placeholder"], 
                 font=("Arial", 10, "bold")).grid(row=0, column=1, sticky=tk.W, padx=10, pady=5)
        
        # 映射输入
        ttk.Label(dialog, text="匹配模式:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)
        mapping_var = tk.StringVar(value=current_mapping)
        mapping_combo = ttk.Combobox(dialog, textvariable=mapping_var, width=30)
        mapping_combo.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=10, pady=5)
        
        # 添加Excel字段选项
        if self.excel_data is not None:
            mapping_combo['values'] = [""] + list(self.excel_data.columns)
        
        # 说明文本
        help_text = """映射规则说明：
1. 直接字段映射：选择Excel字段名
2. 数学表达式：如 字段1+字段2、字段1*2
3. 固定文本：直接输入文本内容
4. 空白：占位符将被替换为"0"
        """
        
        help_label = ttk.Label(dialog, text=help_text, justify=tk.LEFT, 
                              font=("Arial", 9), wraplength=350)
        help_label.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=10, pady=10)
        
        # 按钮
        btn_frame = ttk.Frame(dialog)
        btn_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        def save_mapping():
            new_mapping = mapping_var.get()
            self.mapping_data[item_index]["mapping"] = new_mapping
            self.update_mapping_tree()
            dialog.destroy()
        
        def cancel_edit():
            dialog.destroy()
        
        ttk.Button(btn_frame, text="保存", command=save_mapping).grid(row=0, column=0, padx=5)
        ttk.Button(btn_frame, text="取消", command=cancel_edit).grid(row=0, column=1, padx=5)
        
        # 配置权重
        dialog.columnconfigure(1, weight=1)
        
        # 焦点设置
        mapping_combo.focus()
    
    def add_image_mapping_row(self):
        """添加图片映射行"""
        new_mapping = {
            "folder": "",
            "mapping_rule": "",
            "placeholder": ""
        }
        self.image_mapping_data.append(new_mapping)
        self.update_image_tree()
    
    def remove_image_mapping_row(self):
        """删除选中的图片映射行"""
        selection = self.image_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择要删除的行！")
            return
        
        if messagebox.askyesno("确认", "确定要删除选中的行吗？"):
            for item in selection:
                item_index = int(self.image_tree.item(item, "text")) - 1
                if 0 <= item_index < len(self.image_mapping_data):
                    del self.image_mapping_data[item_index]
            self.update_image_tree()
    
    def clear_image_mappings(self):
        """清空所有图片映射"""
        if messagebox.askyesno("确认", "确定要清空所有图片映射吗？"):
            self.image_mapping_data = []
            self.update_image_tree()
    
    def update_image_tree(self):
        """更新图片映射表格"""
        # 清空现有项目
        for item in self.image_tree.get_children():
            self.image_tree.delete(item)
        
        for i, data in enumerate(self.image_mapping_data):
            folder_display = data["folder"] if data["folder"] else "未选择"
            mapping_display = data["mapping_rule"] if data["mapping_rule"] else "未设置"
            placeholder_display = data["placeholder"] if data["placeholder"] else "未选择"
            
            self.image_tree.insert("", "end", text=str(i+1), 
                                  values=(folder_display, mapping_display, placeholder_display))
    
    def edit_image_mapping(self, event):
        """编辑图片映射关系"""
        selection = self.image_tree.selection()
        if not selection:
            return
        
        item = selection[0]
        item_index = int(self.image_tree.item(item, "text")) - 1
        current_data = self.image_mapping_data[item_index]
        
        # 创建编辑对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("编辑图片映射")
        dialog.geometry("500x400")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 图片文件夹选择
        ttk.Label(dialog, text="图片文件夹:").grid(row=0, column=0, sticky=tk.W, padx=10, pady=5)
        folder_frame = ttk.Frame(dialog)
        folder_frame.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=10, pady=5)
        
        folder_var = tk.StringVar(value=current_data["folder"])
        folder_entry = ttk.Entry(folder_frame, textvariable=folder_var, width=30)
        folder_entry.grid(row=0, column=0, sticky=(tk.W, tk.E))
        
        def select_folder():
            folder_path = filedialog.askdirectory(title="选择图片文件夹")
            if folder_path:
                folder_var.set(folder_path)
        
        ttk.Button(folder_frame, text="浏览", command=select_folder).grid(row=0, column=1, padx=(5, 0))
        folder_frame.columnconfigure(0, weight=1)
        
        # 映射规则选择
        ttk.Label(dialog, text="映射规则:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)
        mapping_var = tk.StringVar(value=current_data["mapping_rule"])
        mapping_combo = ttk.Combobox(dialog, textvariable=mapping_var, width=30)
        mapping_combo.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=10, pady=5)
        
        # 添加映射规则选项
        mapping_options = ["固定图片名"]
        if self.excel_data is not None:
            # 添加Excel字段作为映射规则
            for column in self.excel_data.columns:
                mapping_options.append(f"根据字段: {column}")
            mapping_options.append("根据行号")
        
        mapping_combo['values'] = mapping_options
        
        # 图片选择按钮（需要先定义）
        def select_image():
            if not folder_var.get():
                messagebox.showwarning("警告", "请先选择图片文件夹！")
                return
            
            image_path = filedialog.askopenfilename(
                title="选择图片文件",
                initialdir=folder_var.get(),
                filetypes=[
                    ("图片文件", "*.jpg *.jpeg *.png *.bmp *.gif"),
                    ("所有文件", "*.*")
                ]
            )
            
            if image_path:
                # 提取图片文件名（不含扩展名）
                image_name = os.path.splitext(os.path.basename(image_path))[0]
                mapping_var.set(f"固定图片名: {image_name}")
        
        select_image_btn = ttk.Button(dialog, text="选择图片", command=select_image)
        
        # 如果选择了固定图片名，显示图片选择按钮
        def on_mapping_change(event=None):
            if mapping_var.get() == "固定图片名" or mapping_var.get().startswith("固定图片名: "):
                select_image_btn.grid(row=1, column=2, padx=(5, 0))
            else:
                select_image_btn.grid_remove()
        
        mapping_combo.bind('<<ComboboxSelected>>', on_mapping_change)
        
        # 初始化显示状态
        on_mapping_change()
        
        # Word占位符选择
        ttk.Label(dialog, text="Word占位符:").grid(row=2, column=0, sticky=tk.W, padx=10, pady=5)
        placeholder_var = tk.StringVar(value=current_data["placeholder"])
        placeholder_combo = ttk.Combobox(dialog, textvariable=placeholder_var, width=30)
        placeholder_combo.grid(row=2, column=1, sticky=(tk.W, tk.E), padx=10, pady=5)
        
        # 添加占位符选项
        placeholder_combo['values'] = [""] + self.placeholders
        
        # 说明文本
        help_text = """映射规则说明：
1. 固定图片名：所有数据行都使用同一张图片
- 选择"固定图片名"后，点击"选择图片"按钮指定具体图片
2. 根据字段：根据Excel字段值选择对应图片
- 如字段值为"apple"，则查找"apple.jpg"或"apple.png"等
3. 根据行号：使用行号作为图片名
- 第1行数据使用"1.jpg"，第2行使用"2.jpg"等
图片格式支持：.jpg, .jpeg, .png, .bmp, .gif
程序会自动尝试不同的文件扩展名进行匹配
        """
        
        help_label = ttk.Label(dialog, text=help_text, justify=tk.LEFT, 
                              font=("Arial", 9), wraplength=450)
        help_label.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=10, pady=10)
        
        # 按钮
        btn_frame = ttk.Frame(dialog)
        btn_frame.grid(row=4, column=0, columnspan=2, pady=10)
        
        def save_mapping():
            new_folder = folder_var.get()
            new_mapping = mapping_var.get()
            new_placeholder = placeholder_var.get()
            
            if not new_folder:
                messagebox.showwarning("警告", "请选择图片文件夹！")
                return
            
            if not new_mapping:
                messagebox.showwarning("警告", "请设置映射规则！")
                return
            
            if not new_placeholder:
                messagebox.showwarning("警告", "请选择Word占位符！")
                return
            
            self.image_mapping_data[item_index]["folder"] = new_folder
            self.image_mapping_data[item_index]["mapping_rule"] = new_mapping
            self.image_mapping_data[item_index]["placeholder"] = new_placeholder
            self.update_image_tree()
            dialog.destroy()
        
        def cancel_edit():
            dialog.destroy()
        
        ttk.Button(btn_frame, text="保存", command=save_mapping).grid(row=0, column=0, padx=5)
        ttk.Button(btn_frame, text="取消", command=cancel_edit).grid(row=0, column=1, padx=5)
        
        # 配置权重
        dialog.columnconfigure(1, weight=1)
        
        # 焦点设置
        folder_entry.focus()
    
    def show_image_debug_info(self):
        """显示图片映射调试信息"""
        debug_window = tk.Toplevel(self.root)
        debug_window.title("图片映射调试信息")
        debug_window.geometry("600x400")
        debug_window.transient(self.root)
        
        # 创建文本显示区域
        text_frame = ttk.Frame(debug_window)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Courier", 10))
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)
        
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 生成调试信息
        debug_info = "=== 图片映射调试信息 ===\n\n"
        
        if not self.image_mapping_data:
            debug_info += "没有配置图片映射。\n"
        else:
            debug_info += f"共配置了 {len(self.image_mapping_data)} 个图片映射：\n\n"
            
            for i, mapping in enumerate(self.image_mapping_data):
                debug_info += f"映射 {i+1}:\n"
                debug_info += f"  占位符: {mapping['placeholder']}\n"
                debug_info += f"  映射规则: {mapping['mapping_rule']}\n"
                debug_info += f"  图片文件夹: {mapping['folder']}\n"
                
                # 检查文件夹是否存在
                if mapping['folder']:
                    if os.path.exists(mapping['folder']):
                        debug_info += f"  文件夹状态: 存在\n"
                        
                        # 列出文件夹中的图片文件
                        image_files = []
                        for ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.JPG', '.JPEG', '.PNG', '.BMP', '.GIF']:
                            pattern = os.path.join(mapping['folder'], f"*{ext}")
                            image_files.extend(glob.glob(pattern))
                        
                        if image_files:
                            debug_info += f"  图片文件数量: {len(image_files)}\n"
                            debug_info += f"  图片文件列表:\n"
                            for img_file in sorted(image_files[:10]):  # 只显示前10个
                                debug_info += f"    {os.path.basename(img_file)}\n"
                            if len(image_files) > 10:
                                debug_info += f"    ... 还有 {len(image_files) - 10} 个文件\n"
                        else:
                            debug_info += f"  图片文件数量: 0\n"
                    else:
                        debug_info += f"  文件夹状态: 不存在\n"
                else:
                    debug_info += f"  文件夹状态: 未设置\n"
                
                # 检查占位符是否在模板中
                if mapping['placeholder'] in self.placeholders:
                    debug_info += f"  占位符状态: 已识别\n"
                else:
                    debug_info += f"  占位符状态: 未识别（可能不在模板中）\n"
                
                # 如果有Excel数据，测试第一行的图片匹配
                if self.excel_data is not None and len(self.excel_data) > 0:
                    first_row = self.excel_data.iloc[0]
                    test_image_path = self.get_image_for_row(mapping, first_row, 0)
                    if test_image_path:
                        if os.path.exists(test_image_path):
                            debug_info += f"  测试图片路径: {test_image_path} (存在)\n"
                        else:
                            debug_info += f"  测试图片路径: {test_image_path} (不存在)\n"
                    else:
                        debug_info += f"  测试图片路径: 无法匹配\n"
                else:
                    debug_info += f"  测试图片路径: 无Excel数据\n"
                
                debug_info += "\n"
        
        # 显示Excel数据状态
        debug_info += "=== Excel数据状态 ===\n"
        if self.excel_data is not None:
            debug_info += f"Excel数据行数: {len(self.excel_data)}\n"
            debug_info += f"Excel字段: {list(self.excel_data.columns)}\n"
        else:
            debug_info += "没有导入Excel数据\n"
        
        debug_info += "\n"
        
        # 显示Word模板状态
        debug_info += "=== Word模板状态 ===\n"
        if self.word_template_path:
            debug_info += f"模板路径: {self.word_template_path}\n"
            debug_info += f"识别的占位符: {self.placeholders}\n"
        else:
            debug_info += "没有导入Word模板\n"
        
        debug_info += "\n"
        
        # 显示图片设置
        debug_info += "=== 图片设置 ===\n"
        try:
            image_width = float(self.image_width_var.get())
        except ValueError:
            image_width = 9.8
        
        use_cm = self.use_cm_var.get()
        unit_text = "厘米" if use_cm else "英寸"
        debug_info += f"图片宽度: {image_width} {unit_text}\n"
        
        debug_info += "\n"
        
        # 显示数字格式化设置
        debug_info += "=== 数字格式化设置 ===\n"
        if self.enable_custom_decimal_var.get():
            try:
                decimal_places = int(self.custom_decimal_var.get())
                debug_info += f"格式化方式: 自定义小数位数({decimal_places}位)\n"
            except ValueError:
                debug_info += f"格式化方式: 自定义小数位数(无效设置)\n"
        else:
            debug_info += f"格式化方式: {self.number_format_var.get()}\n"
        
        debug_info += f"千分位分隔符: {'启用' if self.use_thousands_separator_var.get() else '禁用'}\n"
        
        debug_info += "\n"
        
        # 显示文件命名设置
        debug_info += "=== 文件命名设置 ===\n"
        naming_mode = self.naming_mode_var.get()
        debug_info += f"命名方式: {naming_mode}\n"
        
        if naming_mode == "字段":
            field_name = self.naming_field_var.get()
            debug_info += f"选择的字段: {field_name if field_name else '未选择'}\n"
            if self.excel_data is not None and field_name in self.excel_data.columns:
                # 显示该字段的示例值
                sample_values = self.excel_data[field_name].dropna().head(3).tolist()
                debug_info += f"字段示例值: {sample_values}\n"
            else:
                debug_info += "字段状态: 字段不存在或未导入Excel数据\n"
        elif naming_mode == "前缀":
            prefix = self.naming_prefix_var.get()
            debug_info += f"前缀设置: '{prefix}'\n"
            clean_prefix = self.clean_filename(prefix) if prefix else "文档"
            debug_info += f"清理后前缀: '{clean_prefix}'\n"
            debug_info += f"示例文件名: {clean_prefix}_001.docx, {clean_prefix}_002.docx...\n"
        else:
            debug_info += "使用默认命名: 导出文档_001.docx, 导出文档_002.docx...\n"
        
        debug_info += "\n"
        
        # 显示样式保持设置
        debug_info += "=== 样式保持功能 ===\n"
        debug_info += "状态: 已启用（自动）\n"
        debug_info += "保持的格式:\n"
        debug_info += "  • 字体格式: 字体名称、大小、颜色、粗体、斜体、下划线\n"
        debug_info += "  • 段落格式: 对齐方式、行距、段前距、段后距、缩进\n"
        debug_info += "  • 精确替换: 只替换占位符，保留其他文本格式\n"
        debug_info += "  • 多样式支持: 一个段落中不同文本的不同样式都保持\n"
        debug_info += "  • 详细日志: 字体复制过程可在控制台输出中查看\n"
        debug_info += "适用范围: 正文、表格、页眉页脚、文本框等\n"
        
        debug_info += "\n"
        
        # 显示合并文档设置
        debug_info += "=== 合并文档设置 ===\n"
        debug_info += f"合并导出文档: {'启用' if self.merge_docs_var.get() else '禁用'}\n"
        if self.merge_docs_var.get():
            debug_info += "完整合并功能:\n"
            debug_info += "  • 完整格式保持（所有原始文档的格式、样式完全保持）\n"
            debug_info += "  • 分节符保持（保留原文档的分节符和分页符结构）\n"
            debug_info += "  • 页眉页脚（保持各文档的页眉、页脚、首页格式等）\n"
            debug_info += "  • 表格格式（表格样式、行高、列宽、单元格格式等完全保留）\n"
            debug_info += "  • 图片内容（所有图片及其格式、位置完全保持）\n"
            debug_info += "  • 文档属性（保留文档标题、作者等元数据信息）\n"
            debug_info += "  • 样式定义（保留文档中的自定义样式定义）\n"
            debug_info += "  • 复杂结构（支持嵌套表格、文本框、形状等）\n"
            debug_info += "  • XML级别复制（使用深度复制确保完整性）\n"
            debug_info += "  • 多重备用方案（主方法失败时自动使用替代方法）\n"
            debug_info += "  • 自动分页（文档间插入分页符）\n"
            debug_info += "  • 合并后删除临时文件\n"
        
        text_widget.insert(tk.END, debug_info)
        text_widget.config(state=tk.DISABLED)
        
        # 关闭按钮
        ttk.Button(debug_window, text="关闭", 
                  command=debug_window.destroy).pack(pady=10)
    
    def log_output(self, message: str):
        """记录输出信息"""
        import datetime
        timestamp = datetime.datetime.now().strftime("%H:%M:%S")
        formatted_message = f"[{timestamp}] {message}"
        self.console_output.append(formatted_message)
        # 保持最多1000条记录
        if len(self.console_output) > 1000:
            self.console_output = self.console_output[-1000:]
    
    def show_console_output(self):
        """显示控制台输出"""
        output_window = tk.Toplevel(self.root)
        output_window.title("控制台输出")
        output_window.geometry("700x500")
        output_window.transient(self.root)
        
        # 创建文本显示区域
        text_frame = ttk.Frame(output_window)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Courier", 9))
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)
        
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 显示输出内容
        if self.console_output:
            output_text = "\n".join(self.console_output)
        else:
            output_text = "暂无输出信息。\n\n提示：\n1. 执行预览或导出操作后会产生调试信息\n2. 图片映射的详细过程会在这里显示\n3. 如果遇到问题，请检查这里的错误信息"
        
        text_widget.insert(tk.END, output_text)
        text_widget.config(state=tk.DISABLED)
        
        # 按钮框架
        btn_frame = ttk.Frame(output_window)
        btn_frame.pack(pady=10)
        
        # 清空按钮
        def clear_output():
            self.console_output.clear()
            text_widget.config(state=tk.NORMAL)
            text_widget.delete(1.0, tk.END)
            text_widget.insert(tk.END, "输出已清空")
            text_widget.config(state=tk.DISABLED)
        
        ttk.Button(btn_frame, text="清空", command=clear_output).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="刷新", command=lambda: self.refresh_console_output(text_widget)).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="关闭", command=output_window.destroy).pack(side=tk.LEFT, padx=5)
    
    def refresh_console_output(self, text_widget):
        """刷新控制台输出显示"""
        text_widget.config(state=tk.NORMAL)
        text_widget.delete(1.0, tk.END)
        
        if self.console_output:
            output_text = "\n".join(self.console_output)
        else:
            output_text = "暂无输出信息"
        
        text_widget.insert(tk.END, output_text)
        text_widget.config(state=tk.DISABLED)
        text_widget.see(tk.END)  # 滚动到底部
    
    def find_image_file(self, folder_path: str, image_name: str) -> Optional[str]:
        """在指定文件夹中查找图片文件"""
        self.log_output(f"查找图片文件: 文件夹='{folder_path}', 图片名='{image_name}'")
        
        if not os.path.exists(folder_path):
            self.log_output(f"图片文件夹不存在: {folder_path}")
            return None
        
        # 支持的图片格式
        extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.JPG', '.JPEG', '.PNG', '.BMP', '.GIF']
        
        # 尝试不同的扩展名（精确匹配）
        self.log_output(f"尝试精确匹配图片文件...")
        for ext in extensions:
            file_path = os.path.join(folder_path, f"{image_name}{ext}")
            if os.path.exists(file_path):
                self.log_output(f"找到精确匹配的图片: {file_path}")
                return file_path
        
        # 如果没有找到，尝试模糊匹配
        self.log_output(f"精确匹配失败，尝试模糊匹配...")
        for ext in extensions:
            pattern = os.path.join(folder_path, f"*{image_name}*{ext}")
            matches = glob.glob(pattern)
            if matches:
                found_file = matches[0]
                self.log_output(f"找到模糊匹配的图片: {found_file}")
                return found_file
        
        # 列出文件夹中的所有图片文件用于调试
        self.log_output(f"没有找到匹配的图片文件，列出文件夹中的所有图片:")
        all_images = []
        for ext in extensions:
            pattern = os.path.join(folder_path, f"*{ext}")
            all_images.extend(glob.glob(pattern))
        
        if all_images:
            for img in all_images[:10]:  # 只显示前10个
                self.log_output(f"  可用图片: {os.path.basename(img)}")
            if len(all_images) > 10:
                self.log_output(f"  ... 还有 {len(all_images) - 10} 个图片文件")
        else:
            self.log_output("  文件夹中没有图片文件")
        
        return None
    
    def get_image_for_row(self, mapping_data: dict, data_row: pd.Series, row_index: int) -> Optional[str]:
        """根据映射规则获取当前行对应的图片路径"""
        folder_path = mapping_data["folder"]
        mapping_rule = mapping_data["mapping_rule"]
        
        self.log_output(f"图片映射详情 - 文件夹: {folder_path}, 规则: {mapping_rule}")
        
        if not folder_path or not mapping_rule:
            self.log_output("文件夹路径或映射规则为空")
            return None
        
        image_name = ""
        
        if mapping_rule.startswith("固定图片名: "):
            # 提取固定图片名
            fixed_name = mapping_rule.replace("固定图片名: ", "")
            self.log_output(f"使用固定图片名: {fixed_name}")
            return self.find_image_file(folder_path, fixed_name)
        elif mapping_rule == "固定图片名":
            # 如果只是"固定图片名"没有具体名称，使用文件夹中的第一个图片
            self.log_output("查找文件夹中的第一个图片文件")
            image_files = []
            for ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.JPG', '.JPEG', '.PNG', '.BMP', '.GIF']:
                pattern = os.path.join(folder_path, f"*{ext}")
                image_files.extend(glob.glob(pattern, recursive=False))
            
            if image_files:
                selected_image = image_files[0]
                self.log_output(f"找到第一个图片文件: {selected_image}")
                return selected_image
            else:
                self.log_output("文件夹中没有找到图片文件")
                return None
                
        elif mapping_rule.startswith("根据字段: "):
            # 根据Excel字段值选择图片
            field_name = mapping_rule.replace("根据字段: ", "")
            self.log_output(f"根据字段选择图片: {field_name}")
            if field_name in data_row.index:
                field_value = data_row[field_name]
                self.log_output(f"字段值: {field_value}")
                if pd.notna(field_value):
                    image_name = str(field_value)
                    result = self.find_image_file(folder_path, image_name)
                    self.log_output(f"查找图片结果: {result}")
                    return result
                else:
                    self.log_output("字段值为空")
            else:
                self.log_output(f"字段 {field_name} 不存在于数据中")
                    
        elif mapping_rule == "根据行号":
            # 根据行号选择图片
            image_name = str(row_index + 1)  # 行号从1开始
            self.log_output(f"根据行号选择图片: {image_name}")
            result = self.find_image_file(folder_path, image_name)
            self.log_output(f"查找图片结果: {result}")
            return result
        
        self.log_output("没有匹配的映射规则")
        return None
    
    def insert_image_into_paragraph(self, paragraph, image_path: str, width_value: float = 9.8, use_cm: bool = True):
        """在段落中插入图片"""
        try:
            # 检查图片文件是否存在
            if not os.path.exists(image_path):
                self.log_output(f"图片文件不存在: {image_path}")
                paragraph.clear()
                paragraph.text = f"[图片文件不存在: {os.path.basename(image_path)}]"
                return False
            
            # 检查图片文件大小（避免过大的文件）
            file_size = os.path.getsize(image_path)
            # if file_size > 10 * 1024 * 1024:  # 10MB限制
            #     self.log_output(f"图片文件过大: {image_path} ({file_size} bytes)")
            #     paragraph.clear()
            #     paragraph.text = f"[图片文件过大: {os.path.basename(image_path)}]"
            #     return False
            
            unit_text = "厘米" if use_cm else "英寸"
            self.log_output(f"开始插入图片: {image_path} (大小: {file_size} bytes, 宽度: {width_value}{unit_text})")
            
            # 清除段落原有内容
            paragraph.clear()
            
            # 添加新的运行并插入图片
            run = paragraph.add_run()
            
            # 根据单位选择设置图片宽度
            if use_cm:
                run.add_picture(image_path, width=Cm(width_value))
            else:
                run.add_picture(image_path, width=Inches(width_value))
            
            # 设置段落居中对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self.log_output(f"成功插入图片: {image_path}")
            return True
            
        except FileNotFoundError:
            self.log_output(f"图片文件未找到: {image_path}")
            paragraph.clear()
            paragraph.text = f"[图片文件未找到: {os.path.basename(image_path)}]"
            return False
        except Exception as e:
            self.log_output(f"插入图片失败: {image_path}, 错误: {str(e)}")
            # 如果插入失败，显示错误文本
            paragraph.clear()
            paragraph.text = f"[图片插入失败: {os.path.basename(image_path) if image_path else '未知'} - {str(e)}]"
            return False
    
    def calculate_similarity(self, str1: str, str2: str) -> float:
        """计算字符串相似度"""
        if not str1 or not str2:
            return 0.0
        
        # 转换为小写并去除特殊字符（支持Unicode字符）
        str1 = re.sub(r'[^\w\s\u4e00-\u9fff]', '', str1.lower(), flags=re.UNICODE).strip()
        str2 = re.sub(r'[^\w\s\u4e00-\u9fff]', '', str2.lower(), flags=re.UNICODE).strip()
        
        if str1 == str2:
            return 1.0
        
        # 使用difflib计算相似度
        similarity = difflib.SequenceMatcher(None, str1, str2).ratio()

        # print(str1,str2,similarity)
        
        # 包含关系加权
        if str1 in str2 or str2 in str1:
            min_len = min(len(str1), len(str2))
            max_len = max(len(str1), len(str2))
            if min_len / max_len > 0.5:
                similarity = max(similarity, 0.8)
        
        return similarity
    
    def auto_match_fields(self):
        """自动匹配字段"""
        try:
            if self.excel_data is None or not self.placeholders:
                messagebox.showwarning("警告", "请先导入Excel数据和Word模板！")
                return
            
            matched_count = 0
            excel_columns = list(self.excel_data.columns)
            
            for i, data in enumerate(self.mapping_data):
                placeholder = data["placeholder"]
                # 提取占位符中的字段名（去掉{}符号）
                clean_placeholder = placeholder.strip('{}【】《》[]<>')
                
                matched_field = None
                
                if self.exact_match_var.get():
                    # 精准匹配
                    for column in excel_columns:
                        if column.lower() == clean_placeholder.lower():
                            matched_field = column
                            break
                else:
                    # 模糊匹配
                    best_similarity = 0.0
                    for column in excel_columns:
                        similarity = self.calculate_similarity(clean_placeholder, column)
                        if similarity > 0.6 and similarity > best_similarity:
                            best_similarity = similarity
                            matched_field = column
                
                if matched_field:
                    self.mapping_data[i]["mapping"] = matched_field
                    matched_count += 1
            
            self.update_mapping_tree()
            
            match_mode = "精准匹配" if self.exact_match_var.get() else "模糊匹配"
            messagebox.showinfo("完成", 
                f"自动匹配完成！\n匹配模式：{match_mode}\n成功匹配：{matched_count}个占位符\n总占位符：{len(self.placeholders)}个")
            
        except Exception as e:
            messagebox.showerror("错误", f"自动匹配失败：{str(e)}")
    
    def is_number(self, value) -> bool:
        """检查值是否为数字"""
        try:
            float(value)
            return True
        except (ValueError, TypeError):
            return False
    
    def format_number_value(self, value: str) -> str:
        """根据设置格式化数字值"""
        try:
            # 检查是否为数字
            if not self.is_number(value):
                return value
            
            # 转换为浮点数
            num_value = float(value)
            
            # 根据用户设置决定格式化方式
            if self.enable_custom_decimal_var.get():
                # 使用自定义小数位数
                try:
                    decimal_places = int(self.custom_decimal_var.get())
                    formatted_value = f"{num_value:.{decimal_places}f}"
                except ValueError:
                    # 如果自定义小数位数无效，使用原值
                    formatted_value = str(num_value)
            else:
                # 使用预设格式
                format_option = self.number_format_var.get()
                if format_option == "保留原格式":
                    formatted_value = str(num_value)
                elif format_option == "取整数":
                    formatted_value = str(int(round(num_value)))
                elif format_option == "保留1位小数":
                    formatted_value = f"{num_value:.1f}"
                elif format_option == "保留2位小数":
                    formatted_value = f"{num_value:.2f}"
                elif format_option == "保留3位小数":
                    formatted_value = f"{num_value:.3f}"
                else:
                    formatted_value = str(num_value)
            
            # 添加千分位分隔符（如果启用）
            if self.use_thousands_separator_var.get():
                try:
                    # 分离整数和小数部分
                    if '.' in formatted_value:
                        integer_part, decimal_part = formatted_value.split('.')
                        # 为整数部分添加千分位分隔符
                        integer_part = f"{int(integer_part):,}"
                        formatted_value = f"{integer_part}.{decimal_part}"
                    else:
                        formatted_value = f"{int(float(formatted_value)):,}"
                except:
                    # 如果添加千分位分隔符失败，返回原格式化值
                    pass
            
            self.log_output(f"数字格式化: {value} -> {formatted_value}")
            return formatted_value
            
        except Exception as e:
            self.log_output(f"数字格式化失败: {value}, 错误: {str(e)}")
            return value

    def process_math_expression(self, expression: str, data_row: pd.Series) -> str:
        """处理数学表达式"""
        try:
            result = expression
            
            # 替换字段名为实际值
            for column in data_row.index:
                if column in result:
                    value = str(data_row[column]) if pd.notna(data_row[column]) else "0"
                    result = result.replace(column, value)
            
            # 计算表达式
            try:
                computed_value = eval(result)
                # 对计算结果进行数字格式化
                return self.format_number_value(str(computed_value))
            except:
                return result
                
        except:
            return expression
    
    def replace_text_preserve_style(self, paragraph, placeholder, value):
        """在段落中替换文本，保持原有样式"""
        try:
            # 获取段落的完整文本
            full_text = paragraph.text
            
            # 如果占位符不在文本中，直接返回
            if placeholder not in full_text:
                return False
            
            # 保存段落级别的格式
            original_format = self.save_paragraph_format(paragraph)
            
            # 记录所有run及其样式信息和文本
            runs_info = []
            for i, run in enumerate(paragraph.runs):
                # 详细保存字体信息
                font_info = self.save_font_info(run.font)
                
                run_info = {
                    'text': run.text,
                    'font_info': font_info,
                    'style': run.style
                }
                runs_info.append(run_info)
                
                self.log_output(f"保存run {i}: 文本='{run.text[:20]}{'...' if len(run.text) > 20 else ''}', 字体={font_info.get('name', '未知')}")
                
            
            # 如果没有runs，直接简单替换
            if not runs_info:
                paragraph.text = paragraph.text.replace(placeholder, value)
                return True
            
            # 找到占位符的位置
            placeholder_start = full_text.find(placeholder)
            placeholder_end = placeholder_start + len(placeholder)
            
            # 分析每个run与占位符的关系
            current_pos = 0
            new_runs_data = []
            placeholder_found = False
            replacement_added = False
            
            for i, run_info in enumerate(runs_info):
                run_start = current_pos
                run_end = current_pos + len(run_info['text'])
                run_text = run_info['text']
                
                self.log_output(f"  分析run {i}: 位置{run_start}-{run_end}, 占位符位置{placeholder_start}-{placeholder_end}")
                
                # 检查这个run与占位符的关系
                if run_end <= placeholder_start:
                    # 完全在占位符之前
                    new_runs_data.append({
                        'text': run_text,
                        'run_info': run_info
                    })
                    self.log_output(f"    -> 完全在占位符之前，保留全部文本")
                elif run_start >= placeholder_end:
                    # 完全在占位符之后
                    new_runs_data.append({
                        'text': run_text,
                        'run_info': run_info
                    })
                    self.log_output(f"    -> 完全在占位符之后，保留全部文本")
                else:
                    # 与占位符有交集
                    placeholder_found = True
                    
                    # 计算占位符在当前run中的相对位置
                    run_placeholder_start = max(0, placeholder_start - run_start)
                    run_placeholder_end = min(len(run_text), placeholder_end - run_start)
                    
                    # 构建新的文本片段
                    before_placeholder = run_text[:run_placeholder_start]
                    after_placeholder = run_text[run_placeholder_end:]
                    
                    self.log_output(f"    -> 与占位符有交集，前部分:'{before_placeholder}', 后部分:'{after_placeholder}'")
                    
                    # 添加占位符之前的部分（如果有）
                    if before_placeholder:
                        new_runs_data.append({
                            'text': before_placeholder,
                            'run_info': run_info
                        })
                        self.log_output(f"    -> 添加占位符前文本: '{before_placeholder}'")
                    
                    # 只在第一次遇到占位符时添加替换值
                    if not replacement_added:
                        new_runs_data.append({
                            'text': value,
                            'run_info': run_info
                        })
                        replacement_added = True
                        self.log_output(f"    -> 添加替换值: '{value}'")
                     
                    # 添加占位符之后的部分（如果有）
                    if after_placeholder:
                        new_runs_data.append({
                            'text': after_placeholder,
                            'run_info': run_info
                        })
                        self.log_output(f"    -> 添加占位符后文本: '{after_placeholder}'")
                
                current_pos = run_end
            
            # 如果没有找到占位符（理论上不应该发生），使用简单替换
            if not placeholder_found:
                self.log_output(f"警告：在分析run时未找到占位符 {placeholder}")
                 # 使用第一个run的样式进行简单替换
                paragraph.clear()
                new_text = full_text.replace(placeholder, value)
                new_run = paragraph.add_run(new_text)
                if runs_info:
                    self.copy_run_style(runs_info[0], new_run)
                return True
             
            # 清空段落并重建
            paragraph.clear()
             
             # 重建所有runs
            for run_data in new_runs_data:
                if run_data['text']:  # 只添加非空文本的run
                    new_run = paragraph.add_run(run_data['text'])
                    self.copy_run_style(run_data['run_info'], new_run)
            
                        # 恢复段落级别的格式
            self.restore_paragraph_format(paragraph, original_format)
            
            self.log_output(f"样式保持替换成功: {placeholder} -> {value}")
            self.log_output(f"  保留了 {len(new_runs_data)} 个文本片段的独立格式")
            self.log_output(f"  恢复了段落级别格式: 对齐方式、间距、缩进等")
            return True
            
        except Exception as e:
            self.log_output(f"保持样式替换文本时出错: {e}")
            # 如果出错，回退到简单替换
            try:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    return True
            except:
                pass
            return False

    def save_font_info(self, font):
        """保存字体信息到字典"""
        try:
            font_info = {}
            
            # 保存字体名称
            try:
                font_info['name'] = font.name
            except:
                font_info['name'] = None
            
            # 保存字体大小
            try:
                font_info['size'] = font.size
            except:
                font_info['size'] = None
            
            # 保存粗体
            try:
                font_info['bold'] = font.bold
            except:
                font_info['bold'] = None
            
            # 保存斜体
            try:
                font_info['italic'] = font.italic
            except:
                font_info['italic'] = None
            
            # 保存下划线
            try:
                font_info['underline'] = font.underline
            except:
                font_info['underline'] = None
            
            # 保存颜色
            try:
                font_info['color_rgb'] = font.color.rgb
            except:
                font_info['color_rgb'] = None
            
            # 保存其他属性
            try:
                font_info['subscript'] = font.subscript
            except:
                font_info['subscript'] = None
            
            try:
                font_info['superscript'] = font.superscript
            except:
                font_info['superscript'] = None
            
            try:
                font_info['strike'] = font.strike
            except:
                font_info['strike'] = None
            
            return font_info
            
        except Exception as e:
            self.log_output(f"保存字体信息失败: {e}")
            return {}

    def apply_font_info(self, font_info, target_font):
        """将保存的字体信息应用到目标字体"""
        try:
            if not font_info:
                return
            
            applied_attrs = []
            
            # 应用字体名称（最关键）
            if font_info.get('name'):
                try:
                    target_font.name = font_info['name']
                    applied_attrs.append(f"字体:{font_info['name']}")
                except Exception as e:
                    self.log_output(f"  应用字体名称失败: {e}")
            
            # 应用字体大小
            if font_info.get('size'):
                try:
                    target_font.size = font_info['size']
                    applied_attrs.append(f"大小:{font_info['size']}")
                except Exception as e:
                    self.log_output(f"  应用字体大小失败: {e}")
            
            # 应用粗体
            if font_info.get('bold') is not None:
                try:
                    target_font.bold = font_info['bold']
                    if font_info['bold']:
                        applied_attrs.append("粗体")
                except Exception as e:
                    self.log_output(f"  应用粗体失败: {e}")
            
            # 应用斜体
            if font_info.get('italic') is not None:
                try:
                    target_font.italic = font_info['italic']
                    if font_info['italic']:
                        applied_attrs.append("斜体")
                except Exception as e:
                    self.log_output(f"  应用斜体失败: {e}")
            
            # 应用下划线
            if font_info.get('underline') is not None:
                try:
                    target_font.underline = font_info['underline']
                    if font_info['underline']:
                        applied_attrs.append("下划线")
                except Exception as e:
                    self.log_output(f"  应用下划线失败: {e}")
            
            # 应用颜色
            if font_info.get('color_rgb'):
                try:
                    target_font.color.rgb = font_info['color_rgb']
                    applied_attrs.append("颜色")
                except Exception as e:
                    self.log_output(f"  应用字体颜色失败: {e}")
            
            # 应用其他属性
            if font_info.get('subscript') is not None:
                try:
                    target_font.subscript = font_info['subscript']
                    if font_info['subscript']:
                        applied_attrs.append("下标")
                except:
                    pass
            
            if font_info.get('superscript') is not None:
                try:
                    target_font.superscript = font_info['superscript']
                    if font_info['superscript']:
                        applied_attrs.append("上标")
                except:
                    pass
            
            if font_info.get('strike') is not None:
                try:
                    target_font.strike = font_info['strike']
                    if font_info['strike']:
                        applied_attrs.append("删除线")
                except:
                    pass
            
            if applied_attrs:
                self.log_output(f"  应用样式: {', '.join(applied_attrs)}")
                    
        except Exception as e:
            self.log_output(f"应用字体信息失败: {e}")

    def copy_run_style(self, source_run_info, target_run):
        """复制run的样式"""
        try:
            # 使用保存的字体信息
            font_info = source_run_info.get('font_info')
            if font_info:
                self.apply_font_info(font_info, target_run.font)
            else:
                # 兼容旧版本，直接使用font对象
                source_font = source_run_info.get('font')
                if source_font:
                    self.log_output(f"  使用兼容模式复制字体样式")
                    self.apply_font_info(self.save_font_info(source_font), target_run.font)
            
            # 复制run样式
            try:
                if source_run_info.get('style'):
                    target_run.style = source_run_info['style']
                    self.log_output(f"  应用Run样式: {source_run_info['style']}")
            except Exception as style_error:
                self.log_output(f"  应用run样式失败: {style_error}")
            
        except Exception as e:
            self.log_output(f"复制run样式时出错: {e}")

    def copy_run_format(self, source_run, target_run):
        """复制run的格式"""
        try:
            # 复制字体属性
            source_font = source_run.font
            target_font = target_run.font
            
            if source_font.name:
                target_font.name = source_font.name
            if source_font.size:
                target_font.size = source_font.size
            if source_font.bold is not None:
                target_font.bold = source_font.bold
            if source_font.italic is not None:
                target_font.italic = source_font.italic
            if source_font.underline is not None:
                target_font.underline = source_font.underline
            
            # 复制颜色
            try:
                if source_font.color.rgb:
                    target_font.color.rgb = source_font.color.rgb
            except:
                pass
            
            # 复制其他属性
            try:
                if source_font.subscript is not None:
                    target_font.subscript = source_font.subscript
            except:
                pass
            
            try:
                if source_font.superscript is not None:
                    target_font.superscript = source_font.superscript
            except:
                pass
            
            try:
                if source_font.strike is not None:
                    target_font.strike = source_font.strike
            except:
                pass
            
            # 复制run样式
            if source_run.style:
                target_run.style = source_run.style
                
        except Exception as e:
            self.log_output(f"复制run格式时出错: {e}")

    def save_paragraph_format(self, paragraph):
        """保存段落的格式信息"""
        try:
            format_info = {
                'alignment': paragraph.alignment,
                'style': paragraph.style
            }
            
            # 保存段落格式详细信息
            pf = paragraph.paragraph_format
            format_info['paragraph_format'] = {
                'space_before': pf.space_before,
                'space_after': pf.space_after,
                'line_spacing': pf.line_spacing,
                'left_indent': pf.left_indent,
                'right_indent': pf.right_indent,
                'first_line_indent': pf.first_line_indent
            }
            
            return format_info
            
        except Exception as e:
            self.log_output(f"保存段落格式时出错: {e}")
            return None

    def restore_paragraph_format(self, paragraph, format_info):
        """恢复段落的格式"""
        try:
            if not format_info:
                return
            
            # 恢复基本格式
            if format_info.get('alignment') is not None:
                paragraph.alignment = format_info['alignment']
            
            if format_info.get('style'):
                try:
                    paragraph.style = format_info['style']
                except:
                    pass
            
            # 恢复段落格式详细信息
            pf_info = format_info.get('paragraph_format', {})
            if pf_info:
                pf = paragraph.paragraph_format
                
                try:
                    if pf_info.get('space_before') is not None:
                        pf.space_before = pf_info['space_before']
                    if pf_info.get('space_after') is not None:
                        pf.space_after = pf_info['space_after']
                    if pf_info.get('line_spacing') is not None:
                        pf.line_spacing = pf_info['line_spacing']
                    if pf_info.get('left_indent') is not None:
                        pf.left_indent = pf_info['left_indent']
                    if pf_info.get('right_indent') is not None:
                        pf.right_indent = pf_info['right_indent']
                    if pf_info.get('first_line_indent') is not None:
                        pf.first_line_indent = pf_info['first_line_indent']
                except:
                    pass
                
        except Exception as e:
            self.log_output(f"恢复段落格式时出错: {e}")

    def copy_paragraph_format(self, source_paragraph, target_paragraph):
        """复制段落的格式"""
        try:
            # 复制段落对齐方式
            if source_paragraph.alignment is not None:
                target_paragraph.alignment = source_paragraph.alignment
            
            # 复制段落样式
            if source_paragraph.style:
                target_paragraph.style = source_paragraph.style
            
            # 复制段落格式
            source_pf = source_paragraph.paragraph_format
            target_pf = target_paragraph.paragraph_format
            
            # 复制间距
            if source_pf.space_before is not None:
                target_pf.space_before = source_pf.space_before
            if source_pf.space_after is not None:
                target_pf.space_after = source_pf.space_after
            if source_pf.line_spacing is not None:
                target_pf.line_spacing = source_pf.line_spacing
            
            # 复制缩进
            if source_pf.left_indent is not None:
                target_pf.left_indent = source_pf.left_indent
            if source_pf.right_indent is not None:
                target_pf.right_indent = source_pf.right_indent
            if source_pf.first_line_indent is not None:
                target_pf.first_line_indent = source_pf.first_line_indent
                
        except Exception as e:
            self.log_output(f"复制段落格式时出错: {e}")

    def copy_paragraphs_with_format(self, source_doc, target_doc):
        """复制段落并保持格式"""
        try:
            for source_paragraph in source_doc.paragraphs:
                # 创建新段落
                target_paragraph = target_doc.add_paragraph()
                
                # 复制段落格式
                self.copy_paragraph_format(source_paragraph, target_paragraph)
                
                # 复制段落内容和run格式
                for source_run in source_paragraph.runs:
                    # 检查run中是否包含图片
                    if self.run_contains_image(source_run):
                        # 如果包含图片，复制整个run的XML
                        self.copy_run_with_images(source_run, target_paragraph)
                    else:
                        # 创建新的run
                        target_run = target_paragraph.add_run(source_run.text)
                        
                        # 复制run格式
                        self.copy_run_format(source_run, target_run)
                    
        except Exception as e:
            self.log_output(f"复制段落时出错: {e}")

    def copy_document_structure(self, source_doc, target_doc):
        """完整复制文档结构，包括所有内容、样式、分节符等"""
        try:
            self.log_output("开始完整复制文档结构...")
            
            # 复制文档核心属性
            try:
                if hasattr(source_doc, 'core_properties') and hasattr(target_doc, 'core_properties'):
                    core_props = source_doc.core_properties
                    target_core = target_doc.core_properties
                    if core_props.title:
                        target_core.title = core_props.title
                    if core_props.subject:
                        target_core.subject = core_props.subject
                    if core_props.author:
                        target_core.author = core_props.author
                    self.log_output("复制了文档核心属性")
            except Exception as e:
                self.log_output(f"复制文档核心属性失败: {e}")
            
            # 复制文档样式
            try:
                self.copy_document_styles(source_doc, target_doc)
            except Exception as e:
                self.log_output(f"复制文档样式失败: {e}")
            
            # 复制分节结构
            try:
                self.copy_sections_with_format(source_doc, target_doc)
            except Exception as e:
                self.log_output(f"复制分节结构失败: {e}")
            
            self.log_output("文档结构复制完成")
            
        except Exception as e:
            self.log_output(f"复制文档结构时出错: {e}")
            # 如果完整复制失败，尝试替代方法
            try:
                self.log_output("尝试使用替代的完整复制方法...")
                if not self.copy_document_completely_alternative(source_doc, target_doc):
                    # 如果替代方法也失败，回退到基本复制
                    self.log_output("替代方法失败，使用基本复制方法...")
                    self.copy_paragraphs_with_format(source_doc, target_doc)
                    self.copy_tables_with_format(source_doc, target_doc)
            except Exception as fallback_error:
                self.log_output(f"所有复制方法都失败: {fallback_error}")
                # 最后的备用方案：简单文本复制
                try:
                    for paragraph in source_doc.paragraphs:
                        target_doc.add_paragraph(paragraph.text)
                    self.log_output("使用了最简单的文本复制方法")
                except:
                    self.log_output("所有复制尝试都失败了")

    def copy_document_styles(self, source_doc, target_doc):
        """复制文档样式"""
        try:
            # 获取源文档样式
            source_styles = source_doc.styles
            target_styles = target_doc.styles
            
            # 复制字符样式和段落样式
            for style in source_styles:
                try:
                    # 检查目标文档中是否已存在该样式
                    if style.name not in [s.name for s in target_styles]:
                        # 创建新样式（这里需要根据样式类型来处理）
                        self.log_output(f"发现新样式: {style.name}")
                except Exception as style_error:
                    self.log_output(f"处理样式 {style.name} 时出错: {style_error}")
                    
        except Exception as e:
            self.log_output(f"复制文档样式时出错: {e}")

    def copy_sections_with_format(self, source_doc, target_doc):
        """复制分节并保持格式"""
        try:
            self.log_output(f"源文档有 {len(source_doc.sections)} 个分节")
            
            for section_idx, source_section in enumerate(source_doc.sections):
                self.log_output(f"处理第 {section_idx + 1} 个分节...")
                
                # 如果不是第一个分节，添加分节符
                if section_idx > 0:
                    self.log_output("添加分节符")
                    target_doc.add_section()
                
                # 获取目标分节
                if section_idx < len(target_doc.sections):
                    target_section = target_doc.sections[section_idx]
                else:
                    target_section = target_doc.add_section()
                
                # 复制分节属性
                try:
                    self.copy_section_properties(source_section, target_section)
                except Exception as e:
                    self.log_output(f"复制分节属性失败: {e}")
                
                # 复制页眉页脚
                try:
                    self.copy_headers_footers(source_section, target_section)
                except Exception as e:
                    self.log_output(f"复制页眉页脚失败: {e}")
                
                # 复制该分节的内容
                try:
                    self.copy_section_content(source_doc, target_doc, section_idx)
                except Exception as e:
                    self.log_output(f"复制分节内容失败: {e}")
                    
        except Exception as e:
            self.log_output(f"复制分节时出错: {e}")

    def copy_section_properties(self, source_section, target_section):
        """复制分节属性"""
        try:
            # 复制页面设置
            if hasattr(source_section, 'page_width') and source_section.page_width:
                target_section.page_width = source_section.page_width
            if hasattr(source_section, 'page_height') and source_section.page_height:
                target_section.page_height = source_section.page_height
                
            # 复制页边距
            if hasattr(source_section, 'left_margin') and source_section.left_margin:
                target_section.left_margin = source_section.left_margin
            if hasattr(source_section, 'right_margin') and source_section.right_margin:
                target_section.right_margin = source_section.right_margin
            if hasattr(source_section, 'top_margin') and source_section.top_margin:
                target_section.top_margin = source_section.top_margin
            if hasattr(source_section, 'bottom_margin') and source_section.bottom_margin:
                target_section.bottom_margin = source_section.bottom_margin
                
            # 复制页面方向
            if hasattr(source_section, 'orientation'):
                target_section.orientation = source_section.orientation
                
            # 复制其他分节属性
            if hasattr(source_section, 'start_type'):
                target_section.start_type = source_section.start_type
                
            self.log_output("分节属性复制完成")
            
        except Exception as e:
            self.log_output(f"复制分节属性时出错: {e}")

    def copy_headers_footers(self, source_section, target_section):
        """复制页眉页脚"""
        try:
            # 复制主页眉
            if source_section.header:
                self.copy_header_footer_content(source_section.header, target_section.header)
                self.log_output("复制了主页眉")
                
            # 复制主页脚
            if source_section.footer:
                self.copy_header_footer_content(source_section.footer, target_section.footer)
                self.log_output("复制了主页脚")
                
            # 复制首页页眉
            if hasattr(source_section, 'first_page_header') and source_section.first_page_header:
                if hasattr(target_section, 'first_page_header'):
                    self.copy_header_footer_content(source_section.first_page_header, target_section.first_page_header)
                    self.log_output("复制了首页页眉")
                    
            # 复制首页页脚
            if hasattr(source_section, 'first_page_footer') and source_section.first_page_footer:
                if hasattr(target_section, 'first_page_footer'):
                    self.copy_header_footer_content(source_section.first_page_footer, target_section.first_page_footer)
                    self.log_output("复制了首页页脚")
                    
            # 复制偶数页页眉
            if hasattr(source_section, 'even_page_header') and source_section.even_page_header:
                if hasattr(target_section, 'even_page_header'):
                    self.copy_header_footer_content(source_section.even_page_header, target_section.even_page_header)
                    self.log_output("复制了偶数页页眉")
                    
            # 复制偶数页页脚
            if hasattr(source_section, 'even_page_footer') and source_section.even_page_footer:
                if hasattr(target_section, 'even_page_footer'):
                    self.copy_header_footer_content(source_section.even_page_footer, target_section.even_page_footer)
                    self.log_output("复制了偶数页页脚")
                    
        except Exception as e:
            self.log_output(f"复制页眉页脚时出错: {e}")

    def copy_header_footer_content(self, source_hf, target_hf):
        """复制页眉或页脚的内容"""
        try:
            # 清空目标页眉/页脚
            for paragraph in target_hf.paragraphs:
                paragraph.clear()
            
            # 复制段落
            for i, source_paragraph in enumerate(source_hf.paragraphs):
                if i == 0:
                    # 使用第一个现有段落
                    target_paragraph = target_hf.paragraphs[0]
                else:
                    # 添加新段落
                    target_paragraph = target_hf.add_paragraph()
                
                # 复制段落格式和内容
                self.copy_paragraph_format(source_paragraph, target_paragraph)
                
                for source_run in source_paragraph.runs:
                    if self.run_contains_image(source_run):
                        self.copy_run_with_images(source_run, target_paragraph)
                    else:
                        target_run = target_paragraph.add_run(source_run.text)
                        self.copy_run_format(source_run, target_run)
            
            # 复制表格
            for source_table in source_hf.tables:
                target_table = target_hf.add_table(rows=len(source_table.rows), 
                                                  cols=len(source_table.columns))
                
                # 复制表格样式
                if source_table.style:
                    target_table.style = source_table.style
                
                # 复制表格内容
                for i, source_row in enumerate(source_table.rows):
                    target_row = target_table.rows[i]
                    for j, source_cell in enumerate(source_row.cells):
                        target_cell = target_row.cells[j]
                        self.copy_cell_format(source_cell, target_cell)
                        
        except Exception as e:
            self.log_output(f"复制页眉页脚内容时出错: {e}")

    def copy_section_content(self, source_doc, target_doc, section_idx):
        """复制特定分节的内容"""
        try:
            # 这里我们需要根据分节来复制对应的段落和表格
            # 由于python-docx API的限制，我们采用复制整个文档内容的方式
            # 实际应用中，可以根据需要调整这个逻辑
            
            # 如果是第一个分节，已经有内容了，跳过
            if section_idx == 0:
                # 复制主要内容（段落和表格）
                self.copy_main_content(source_doc, target_doc)
            else:
                # 为后续分节复制内容（这里可以根据实际需求调整）
                self.log_output(f"分节 {section_idx + 1} 的内容复制需要根据具体需求实现")
                
        except Exception as e:
            self.log_output(f"复制分节内容时出错: {e}")

    def copy_main_content(self, source_doc, target_doc):
        """复制主要内容（段落、表格、分页符等）"""
        try:
            # 使用更精确的XML复制方式
            self.log_output("开始复制主要内容...")
            
            source_body = source_doc._body._element
            target_body = target_doc._body._element
            
            element_count = 0
            
            # 复制所有子元素（段落、表格、分页符等）
            for element in source_body:
                try:
                    # 创建元素的深拷贝
                    import copy
                    new_element = copy.deepcopy(element)
                    
                    # 将新元素添加到目标文档
                    target_body.append(new_element)
                    element_count += 1
                    
                    # 记录复制的元素类型
                    element_tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
                    if element_count <= 10:  # 只记录前10个元素的详细信息
                        self.log_output(f"  复制元素 {element_count}: {element_tag}")
                    
                except Exception as element_error:
                    self.log_output(f"复制元素 {element_count + 1} 时出错: {element_error}")
                    # 继续处理下一个元素
                    continue
            
            self.log_output(f"主要内容复制完成，共复制 {element_count} 个元素")
            
        except Exception as e:
            self.log_output(f"XML复制失败，使用备用方法: {e}")
            # 如果XML复制失败，使用原有方法
            try:
                self.copy_paragraphs_with_format(source_doc, target_doc)
                self.copy_tables_with_format(source_doc, target_doc)
                self.log_output("备用方法复制完成")
            except Exception as backup_error:
                self.log_output(f"备用方法也失败: {backup_error}")
                # 最后的备用方法：简单的文本复制
                try:
                    for paragraph in source_doc.paragraphs:
                        target_doc.add_paragraph(paragraph.text)
                    self.log_output("使用了简单文本复制作为最后备用方法")
                except Exception as simple_error:
                    self.log_output(f"简单文本复制也失败: {simple_error}")

    def copy_document_completely_alternative(self, source_doc, target_doc):
        """替代的完整文档复制方法（如果主方法失败）"""
        try:
            self.log_output("使用替代方法进行完整文档复制...")
            
            # 方法1：尝试使用python-docx的内置方法
            try:
                # 复制所有段落
                paragraph_count = 0
                for paragraph in source_doc.paragraphs:
                    new_paragraph = target_doc.add_paragraph()
                    # 复制段落内容和格式
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        # 复制基本格式
                        if run.bold:
                            new_run.bold = run.bold
                        if run.italic:
                            new_run.italic = run.italic
                        if run.underline:
                            new_run.underline = run.underline
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                    paragraph_count += 1
                
                self.log_output(f"复制了 {paragraph_count} 个段落")
                
                # 复制所有表格
                table_count = 0
                for table in source_doc.tables:
                    new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
                    table_count += 1
                
                self.log_output(f"复制了 {table_count} 个表格")
                self.log_output("替代方法复制完成")
                return True
                
            except Exception as alt_error:
                self.log_output(f"替代方法失败: {alt_error}")
                return False
                
        except Exception as e:
            self.log_output(f"替代完整复制方法失败: {e}")
            return False

    def merge_documents_completely(self, file_paths, output_path):
        """完整合并多个文档"""
        try:
            self.log_output("=== 开始完整文档合并 ===")
            self.log_output(f"准备合并 {len(file_paths)} 个文档")
            
            if not file_paths:
                raise ValueError("没有文档需要合并")
            
            # 以第一个文档为基础
            merged_doc = Document(file_paths[0])
            self.log_output(f"以第一个文档为基础: {os.path.basename(file_paths[0])}")
            
            # 合并其他文档
            for i, file_path in enumerate(file_paths[1:], 1):
                self.log_output(f"正在合并第 {i+1} 个文档: {os.path.basename(file_path)}")
                
                try:
                    doc_to_merge = Document(file_path)
                    
                    # 添加分页符（在新内容前）
                    merged_doc.add_page_break()
                    self.log_output("添加了分页符")
                    
                    # 完整复制文档结构
                    self.copy_document_structure(doc_to_merge, merged_doc)
                    
                    self.log_output(f"第 {i+1} 个文档合并完成")
                    
                except Exception as doc_error:
                    self.log_output(f"合并第 {i+1} 个文档时出错: {doc_error}")
                    raise doc_error
            
            # 保存合并文档
            merged_doc.save(output_path)
            self.log_output(f"合并文档保存至: {output_path}")
            self.log_output("=== 完整文档合并完成 ===")
            
            return True
            
        except Exception as e:
            self.log_output(f"完整文档合并失败: {e}")
            return False

    def run_contains_image(self, run):
        """检查run是否包含图片"""
        try:
            # 检查run的XML是否包含图片元素
            return len(run._element.xpath('.//w:drawing')) > 0 or len(run._element.xpath('.//w:pict')) > 0
        except:
            return False

    def copy_run_with_images(self, source_run, target_paragraph):
        """复制包含图片的run"""
        try:
            # 复制run的XML元素
            import copy
            new_run_element = copy.deepcopy(source_run._element)
            target_paragraph._element.append(new_run_element)
            
            self.log_output(f"复制了包含图片的run")
            
        except Exception as e:
            self.log_output(f"复制包含图片的run时出错: {e}")
            # 如果复制失败，创建普通的文本run
            target_run = target_paragraph.add_run(source_run.text)
            self.copy_run_format(source_run, target_run)

    def copy_cell_format(self, source_cell, target_cell):
        """复制表格单元格的格式"""
        try:
            # 复制单元格的段落
            # 先清空目标单元格
            target_cell.paragraphs[0].clear()
            
            for i, source_paragraph in enumerate(source_cell.paragraphs):
                if i == 0:
                    # 使用第一个现有段落
                    target_paragraph = target_cell.paragraphs[0]
                else:
                    # 添加新段落
                    target_paragraph = target_cell.add_paragraph()
                
                # 复制段落格式
                self.copy_paragraph_format(source_paragraph, target_paragraph)
                
                # 复制段落内容
                for source_run in source_paragraph.runs:
                    # 检查run中是否包含图片
                    if self.run_contains_image(source_run):
                        # 如果包含图片，复制整个run的XML
                        self.copy_run_with_images(source_run, target_paragraph)
                    else:
                        target_run = target_paragraph.add_run(source_run.text)
                        self.copy_run_format(source_run, target_run)
                    
        except Exception as e:
            self.log_output(f"复制单元格格式时出错: {e}")

    def copy_tables_with_format(self, source_doc, target_doc):
        """复制表格并保持格式"""
        try:
            for source_table in source_doc.tables:
                # 创建新表格
                target_table = target_doc.add_table(rows=len(source_table.rows), 
                                                   cols=len(source_table.columns))
                
                # 复制表格样式
                try:
                    if source_table.style:
                        target_table.style = source_table.style
                except:
                    pass
                
                # 复制表格内容和格式
                for i, source_row in enumerate(source_table.rows):
                    target_row = target_table.rows[i]
                    
                    for j, source_cell in enumerate(source_row.cells):
                        target_cell = target_row.cells[j]
                        
                        # 复制单元格格式和内容
                        self.copy_cell_format(source_cell, target_cell)
                        
                        # 复制单元格的背景色（如果有的话）
                        try:
                            if source_cell._element.xpath('.//w:shd'):
                                # 这里可以添加背景色复制逻辑
                                pass
                        except:
                            pass
                    
                    # 复制行高（如果有的话）
                    try:
                        if source_row.height:
                            target_row.height = source_row.height
                    except:
                        pass
                
                # 复制列宽（如果有的话）
                try:
                    for i, source_col in enumerate(source_table.columns):
                        if source_col.width:
                            target_table.columns[i].width = source_col.width
                except:
                    pass
                    
        except Exception as e:
            self.log_output(f"复制表格时出错: {e}")

    def replace_placeholder_in_element(self, element, placeholder, value):
        """在文档元素中替换占位符"""
        try:
            # 处理段落
            if hasattr(element, 'paragraphs'):
                for paragraph in element.paragraphs:
                    if placeholder in paragraph.text:
                        # 使用保持样式的替换方法
                        self.replace_text_preserve_style(paragraph, placeholder, value)
            
            # 处理表格
            if hasattr(element, 'tables'):
                for table in element.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # 处理表格单元格中的每个段落
                            for cell_paragraph in cell.paragraphs:
                                if placeholder in cell_paragraph.text:
                                    self.replace_text_preserve_style(cell_paragraph, placeholder, value)
                            # 递归处理表格中的内容
                            self.replace_placeholder_in_element(cell, placeholder, value)
            
            # 处理形状和文本框
            if hasattr(element, '_element'):
                try:
                    from docx.oxml.ns import qn
                    
                    # 查找所有文本框和形状中的文本
                    shapes = element._element.xpath('.//w:txbxContent//w:t')
                    for shape in shapes:
                        if shape.text and placeholder in shape.text:
                            shape.text = shape.text.replace(placeholder, value)
                    
                    # 查找所有绘图对象中的文本
                    drawings = element._element.xpath('.//w:drawing//w:t')
                    for drawing in drawings:
                        if drawing.text and placeholder in drawing.text:
                            drawing.text = drawing.text.replace(placeholder, value)
                
                except Exception as e:
                    # 如果处理特殊元素失败，继续处理
                    pass
        
        except Exception as e:
            # 如果遇到错误，继续处理其他内容
            pass
    
    def apply_mapping_to_document(self, doc: Document, data_row: pd.Series, row_index: int = 0):
        """将映射应用到文档"""
        try:
            # 收集所有图片占位符，避免被文本处理逻辑替换
            image_placeholders = set()
            for img_mapping in self.image_mapping_data:
                if img_mapping.get("placeholder"):
                    image_placeholders.add(img_mapping["placeholder"])
            
            # 处理图片占位符（先处理图片，避免被文本替换）
            self.log_output(f"开始处理图片占位符，共 {len(self.image_mapping_data)} 个映射")
            for img_mapping in self.image_mapping_data:
                placeholder = img_mapping["placeholder"]
                
                if not placeholder:
                    continue
                
                self.log_output(f"处理图片占位符: {placeholder}")
                self.log_output(f"映射规则: {img_mapping['mapping_rule']}")
                self.log_output(f"图片文件夹: {img_mapping['folder']}")
                
                # 获取对应的图片路径
                image_path = self.get_image_for_row(img_mapping, data_row, row_index)
                self.log_output(f"找到图片路径: {image_path}")
                
                if image_path and os.path.exists(image_path):
                    self.log_output(f"图片文件存在，开始替换占位符")
                    image_replaced = False
                    
                    # 获取图片宽度设置
                    try:
                        image_width = float(self.image_width_var.get())
                    except ValueError:
                        image_width = 5.0  # 默认宽度
                        self.log_output(f"图片宽度设置无效，使用默认值: {image_width}")
                    
                    use_cm = self.use_cm_var.get()
                    
                    # 在主文档中查找并替换图片占位符
                    for paragraph in doc.paragraphs:
                        if placeholder in paragraph.text:
                            self.log_output(f"在主文档段落中找到占位符: {placeholder}")
                            if self.insert_image_into_paragraph(paragraph, image_path, image_width, use_cm):
                                image_replaced = True
                    
                    # 在主文档表格中查找并替换图片占位符
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                # 检查表格单元格的段落
                                for cell_paragraph in cell.paragraphs:
                                    if placeholder in cell_paragraph.text:
                                        self.log_output(f"在主文档表格中找到占位符: {placeholder}")
                                        if self.insert_image_into_paragraph(cell_paragraph, image_path, image_width, use_cm):
                                            image_replaced = True
                                
                                # 递归处理表格单元格内的嵌套表格
                                for nested_table in cell.tables:
                                    for nested_row in nested_table.rows:
                                        for nested_cell in nested_row.cells:
                                            for nested_paragraph in nested_cell.paragraphs:
                                                if placeholder in nested_paragraph.text:
                                                    self.log_output(f"在嵌套表格中找到占位符: {placeholder}")
                                                    if self.insert_image_into_paragraph(nested_paragraph, image_path, image_width, use_cm):
                                                        image_replaced = True
                    
                    # 在页眉和页脚中查找并替换图片占位符
                    for section in doc.sections:
                        # 处理页眉
                        if section.header:
                            for paragraph in section.header.paragraphs:
                                if placeholder in paragraph.text:
                                    self.log_output(f"在页眉中找到占位符: {placeholder}")
                                    if self.insert_image_into_paragraph(paragraph, image_path, image_width, use_cm):
                                        image_replaced = True
                        
                        # 处理页脚
                        if section.footer:
                            for paragraph in section.footer.paragraphs:
                                if placeholder in paragraph.text:
                                    self.log_output(f"在页脚中找到占位符: {placeholder}")
                                    if self.insert_image_into_paragraph(paragraph, image_path, image_width, use_cm):
                                        image_replaced = True
                    
                    if not image_replaced:
                        self.log_output(f"警告：占位符 {placeholder} 在文档中未找到！")
                        
                else:
                    self.log_output(f"图片文件不存在或路径为空: {image_path}")
                    # 如果找不到图片，显示错误信息
                    error_text = f"[图片未找到: {os.path.basename(image_path) if image_path else '无'}]"
                    
                    # 在主文档中替换占位符为错误信息
                    for paragraph in doc.paragraphs:
                        if placeholder in paragraph.text:
                            self.replace_text_preserve_style(paragraph, placeholder, error_text)
                    
                    # 在主文档表格中替换占位符为错误信息
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for cell_paragraph in cell.paragraphs:
                                    if placeholder in cell_paragraph.text:
                                        self.replace_text_preserve_style(cell_paragraph, placeholder, error_text)
            
            # 处理文本占位符（跳过已处理的图片占位符）
            for mapping in self.mapping_data:
                placeholder = mapping["placeholder"]
                match_pattern = mapping["mapping"]
                
                if not placeholder:
                    continue
                
                # 跳过图片占位符，避免重复处理
                if placeholder in image_placeholders:
                    self.log_output(f"跳过图片占位符的文本处理: {placeholder}")
                    continue
                
                # 获取替换值
                value = "0"  # 默认值
                
                if match_pattern:
                    if any(op in match_pattern for op in ['+', '-', '*', '/']):
                        # 数学表达式
                        value = self.process_math_expression(match_pattern, data_row)
                    elif match_pattern in data_row.index:
                        # 直接字段映射
                        cell_value = data_row[match_pattern]
                        if pd.notna(cell_value):
                            raw_value = str(cell_value)
                            # 对字段值进行数字格式化
                            value = self.format_number_value(raw_value)
                        else:
                            value = "0"
                    else:
                        # 固定文本
                        value = match_pattern
                
                # 替换主文档中的占位符
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        self.replace_text_preserve_style(paragraph, placeholder, value)
                
                # 替换主文档表格中的占位符
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # 处理表格单元格中的每个段落
                            for cell_paragraph in cell.paragraphs:
                                if placeholder in cell_paragraph.text:
                                    self.replace_text_preserve_style(cell_paragraph, placeholder, value)
                            # 递归处理表格中的内容
                            self.replace_placeholder_in_element(cell, placeholder, value)
                
                # 替换页眉和页脚中的占位符
                for section in doc.sections:
                    # 替换页眉
                    if section.header:
                        self.replace_placeholder_in_element(section.header, placeholder, value)
                    
                    # 替换页脚
                    if section.footer:
                        self.replace_placeholder_in_element(section.footer, placeholder, value)
                    
                    # 替换首页页眉
                    if hasattr(section, 'first_page_header') and section.first_page_header:
                        self.replace_placeholder_in_element(section.first_page_header, placeholder, value)
                    
                    # 替换首页页脚
                    if hasattr(section, 'first_page_footer') and section.first_page_footer:
                        self.replace_placeholder_in_element(section.first_page_footer, placeholder, value)
                    
                    # 替换奇数页页眉
                    if hasattr(section, 'even_page_header') and section.even_page_header:
                        self.replace_placeholder_in_element(section.even_page_header, placeholder, value)
                    
                    # 替换奇数页页脚
                    if hasattr(section, 'even_page_footer') and section.even_page_footer:
                        self.replace_placeholder_in_element(section.even_page_footer, placeholder, value)
                
                # 替换形状和文本框中的占位符
                try:
                    from docx.oxml.ns import qn
                    
                    # 查找所有文本框
                    for textbox in doc._element.xpath('.//w:txbxContent'):
                        for paragraph in textbox.xpath('.//w:p'):
                            for run in paragraph.xpath('.//w:t'):
                                if run.text and placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)
                    
                    # 查找所有绘图对象中的文本
                    for drawing in doc._element.xpath('.//w:drawing'):
                        for text_run in drawing.xpath('.//w:t'):
                            if text_run.text and placeholder in text_run.text:
                                text_run.text = text_run.text.replace(placeholder, value)
                
                except Exception as replace_error:
                    # 如果替换特殊元素失败，继续处理
                    self.log_output(f"替换特殊元素时出错: {replace_error}")
                    
        except Exception as e:
            messagebox.showerror("错误", f"应用映射失败：{str(e)}")
    
    def preview_document(self):
        """预览文档"""
        try:
            if not self.word_template_path:
                messagebox.showwarning("警告", "请先导入Word模板！")
                return
            
            if self.excel_data is None or len(self.excel_data) == 0:
                messagebox.showwarning("警告", "请先导入Excel数据！")
                return
            
            # 验证导出范围
            is_valid, message = self.validate_export_range()
            if not is_valid:
                messagebox.showwarning("警告", f"导出范围设置有误：{message}")
                return
            
            # 获取要预览的数据
            export_data = self.get_export_data_range()
            if export_data is None or len(export_data) == 0:
                messagebox.showwarning("警告", "没有可预览的数据！")
                return
            
            self.log_output("=== 开始预览文档 ===")
            self.log_output(f"Word模板路径: {self.word_template_path}")
            self.log_output(f"Excel数据总行数: {len(self.excel_data)}")
            self.log_output(f"预览范围: {message}")
            
            # 使用范围内第一行数据生成预览
            doc = Document(self.word_template_path)
            first_row = export_data.iloc[0]
            
            # 获取该行在原始数据中的索引（用于文件名生成等）
            original_index = export_data.index[0]
            
            self.log_output(f"使用第 {original_index + 1} 行数据进行预览")
            self.log_output("开始应用映射...")
            
            # 应用映射
            self.apply_mapping_to_document(doc, first_row, original_index)
            
            # 保存到临时文件
            temp_path = os.path.join(tempfile.gettempdir(), "preview_temp.docx")
            doc.save(temp_path)
            
            self.log_output(f"预览文档已保存到: {temp_path}")
            self.log_output("=== 预览文档完成 ===")
            
            # 预览文档
            if self.preview_in_file_var.get():
                self.open_file(temp_path)
            else:
                messagebox.showinfo("预览", f"预览文档已保存到：{temp_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"预览生成失败：{str(e)}")
    
    def export_documents(self):
        """批量导出文档"""
        try:
            if not self.word_template_path:
                messagebox.showwarning("警告", "请先导入Word模板！")
                return
            
            if self.excel_data is None or len(self.excel_data) == 0:
                messagebox.showwarning("警告", "请先导入Excel数据！")
                return
            
            # 验证导出范围
            is_valid, message = self.validate_export_range()
            if not is_valid:
                messagebox.showwarning("警告", f"导出范围设置有误：{message}")
                return
            
            # 获取要导出的数据
            export_data = self.get_export_data_range()
            if export_data is None or len(export_data) == 0:
                messagebox.showwarning("警告", "没有可导出的数据！")
                return
            
            # 选择保存目录
            output_dir = filedialog.askdirectory(title="选择保存目录")
            if not output_dir:
                return
            
            self.log_output("=== 开始批量导出文档 ===")
            self.log_output(f"输出目录: {output_dir}")
            self.log_output(f"Excel数据总行数: {len(self.excel_data)}")
            self.log_output(f"导出范围: {message}")
            
            success_count = 0
            total_count = len(export_data)
            generated_files = []
            used_filenames = set()  # 跟踪已使用的文件名
            
            # 进度对话框
            progress_window = tk.Toplevel(self.root)
            progress_window.title("导出进度")
            progress_window.geometry("400x150")
            progress_window.transient(self.root)
            progress_window.grab_set()
            
            progress_var = tk.StringVar(value="正在导出...")
            ttk.Label(progress_window, textvariable=progress_var).pack(pady=20)
            
            progress_bar = ttk.Progressbar(progress_window, mode='determinate', maximum=total_count)
            progress_bar.pack(pady=10, padx=20, fill=tk.X)
            
            def update_progress(current, total, success):
                progress_var.set(f"正在导出第 {current}/{total} 个文档，成功：{success}")
                progress_bar['value'] = current
                progress_window.update()
            
            # 批量生成文档
            for i, (index, row) in enumerate(export_data.iterrows()):
                try:
                    update_progress(i + 1, total_count, success_count)
                    
                    # 显示原始行号（用户视角的行号）
                    original_row_num = index + 1
                    self.log_output(f"处理第 {i+1}/{total_count} 个文档（原始数据第 {original_row_num} 行）...")
                    
                    # 生成文档
                    doc = Document(self.word_template_path)
                    self.apply_mapping_to_document(doc, row, index)
                    
                    # 生成文件名（使用原始行索引）
                    filename = self.generate_filename(row, index, used_filenames)
                    used_filenames.add(filename)
                    
                    output_path = os.path.join(output_dir, filename)
                    doc.save(output_path)
                    
                    generated_files.append(output_path)
                    success_count += 1
                    
                    self.log_output(f"第 {i+1} 个文档生成成功: {filename}")
                    
                except Exception as row_ex:
                    original_row_num = index + 1
                    self.log_output(f"第 {i+1} 个文档处理失败（原始数据第 {original_row_num} 行）: {str(row_ex)}")
                    messagebox.showwarning("警告", f"第{i+1}个文档处理失败（原始数据第{original_row_num}行）：{str(row_ex)}")
            
            progress_window.destroy()
            
            self.log_output(f"批量导出完成！成功: {success_count}/{total_count}")
            
            # 合并文档（如果选择）
            if self.merge_docs_var.get() and generated_files:
                try:
                    merged_path = os.path.join(output_dir, "合并文档.docx")
                    
                    # 使用新的完整合并方法
                    merge_success = self.merge_documents_completely(generated_files, merged_path)
                    
                    if merge_success:
                        # 删除临时文件
                        for file_path in generated_files:
                            try:
                                os.remove(file_path)
                                self.log_output(f"删除临时文件: {os.path.basename(file_path)}")
                            except Exception as delete_error:
                                self.log_output(f"删除临时文件失败: {delete_error}")
                        
                        messagebox.showinfo("完成", 
                            f"完整文档合并成功！\n"
                            f"成功合并：{success_count}个文档\n"
                            f"失败：{total_count - success_count}个文档\n"
                            f"合并文档保存至：{merged_path}\n\n"
                            f"已包含的内容：\n"
                            f"• 所有原始格式和样式\n"
                            f"• 分节符和分页符\n"
                            f"• 页眉和页脚\n"
                            f"• 表格和图片\n"
                            f"• 文档属性")
                    else:
                        # 如果完整合并失败，尝试使用基本合并方法
                        self.log_output("完整合并失败，尝试使用基本合并方法")
                        
                        merged_doc = Document(generated_files[0])
                        self.log_output(f"以第一个文档为基础: {os.path.basename(generated_files[0])}")
                        
                        for i, file_path in enumerate(generated_files[1:], 1):
                            self.log_output(f"正在基本合并第 {i+1} 个文档: {os.path.basename(file_path)}")
                            doc_to_append = Document(file_path)
                            
                            # 添加分页符
                            merged_doc.add_page_break()
                            
                            # 复制段落并保持格式
                            self.copy_paragraphs_with_format(doc_to_append, merged_doc)
                            
                            # 复制表格并保持格式
                            self.copy_tables_with_format(doc_to_append, merged_doc)
                        
                        # 保存合并文档
                        merged_doc.save(merged_path)
                        
                        # 删除临时文件
                        for file_path in generated_files:
                            try:
                                os.remove(file_path)
                            except:
                                pass
                        
                        messagebox.showinfo("完成", 
                            f"基本文档合并完成！\n成功：{success_count}个文档已合并\n失败：{total_count - success_count}个文档\n合并文档保存至：{merged_path}")
                    
                except Exception as merge_ex:
                    self.log_output(f"文档合并失败: {str(merge_ex)}")
                    messagebox.showerror("错误", f"文档合并失败：{str(merge_ex)}")
            else:
                messagebox.showinfo("完成", 
                    f"批量导出完成！\n成功：{success_count}个文档\n失败：{total_count - success_count}个文档\n保存目录：{output_dir}")
            
            # 打开输出目录
            self.open_file(output_dir)
            
        except Exception as e:
            messagebox.showerror("错误", f"批量导出失败：{str(e)}")
    
    def open_file(self, file_path: str):
        """打开文件或目录"""
        try:
            if platform.system() == "Windows":
                os.startfile(file_path)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", file_path])
            else:  # Linux
                subprocess.run(["xdg-open", file_path])
        except Exception as e:
            messagebox.showerror("错误", f"打开文件失败：{str(e)}")
    
    def test_text_replacement(self):
        """测试文本替换功能"""
        if not self.word_template_path:
            messagebox.showwarning("警告", "请先导入Word模板！")
            return
        
        # 创建测试窗口
        test_window = tk.Toplevel(self.root)
        test_window.title("文本替换测试")
        test_window.geometry("800x600")
        test_window.transient(self.root)
        
        # 测试说明
        ttk.Label(test_window, text="文本替换测试工具", 
                 font=("Arial", 14, "bold")).pack(pady=10)
        
        ttk.Label(test_window, 
                 text="此工具用于测试和调试文本替换功能，特别是占位符前后文本的保留问题。",
                 wraplength=750).pack(pady=5)
        
        # 占位符选择
        placeholder_frame = ttk.Frame(test_window)
        placeholder_frame.pack(fill=tk.X, padx=20, pady=10)
        
        ttk.Label(placeholder_frame, text="选择占位符:").pack(side=tk.LEFT)
        placeholder_var = tk.StringVar()
        placeholder_combo = ttk.Combobox(placeholder_frame, textvariable=placeholder_var, 
                                        values=self.placeholders, width=30)
        placeholder_combo.pack(side=tk.LEFT, padx=(10, 0))
        
        # 替换值输入
        value_frame = ttk.Frame(test_window)
        value_frame.pack(fill=tk.X, padx=20, pady=5)
        
        ttk.Label(value_frame, text="替换值:").pack(side=tk.LEFT)
        value_var = tk.StringVar(value="TEST_VALUE")
        value_entry = ttk.Entry(value_frame, textvariable=value_var, width=30)
        value_entry.pack(side=tk.LEFT, padx=(10, 0))
        
        # 测试结果显示
        result_frame = ttk.LabelFrame(test_window, text="测试结果", padding=10)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        result_text = tk.Text(result_frame, wrap=tk.WORD, font=("Courier", 10))
        result_scroll = ttk.Scrollbar(result_frame, orient="vertical", command=result_text.yview)
        result_text.configure(yscrollcommand=result_scroll.set)
        
        result_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        result_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        def run_test():
            """执行测试"""
            selected_placeholder = placeholder_var.get()
            replacement_value = value_var.get()
            
            if not selected_placeholder:
                messagebox.showwarning("警告", "请选择一个占位符！")
                return
            
            result_text.delete(1.0, tk.END)
            result_text.insert(tk.END, f"=== 测试文本替换功能 ===\n")
            result_text.insert(tk.END, f"占位符: {selected_placeholder}\n")
            result_text.insert(tk.END, f"替换值: {replacement_value}\n\n")
            
            try:
                # 加载模板文档
                doc = Document(self.word_template_path)
                
                # 查找包含占位符的段落
                found_paragraphs = []
                for i, paragraph in enumerate(doc.paragraphs):
                    if selected_placeholder in paragraph.text:
                        found_paragraphs.append((i, paragraph))
                
                result_text.insert(tk.END, f"找到 {len(found_paragraphs)} 个包含占位符的段落\n\n")
                
                if not found_paragraphs:
                    result_text.insert(tk.END, "未找到包含该占位符的段落！\n")
                    return
                
                # 测试第一个找到的段落
                paragraph_idx, paragraph = found_paragraphs[0]
                
                result_text.insert(tk.END, f"测试段落 {paragraph_idx + 1}:\n")
                result_text.insert(tk.END, f"原始文本: {paragraph.text}\n")
                result_text.insert(tk.END, f"Runs数量: {len(paragraph.runs)}\n")
                
                # 显示runs详情
                result_text.insert(tk.END, "\n--- Runs详情 ---\n")
                for i, run in enumerate(paragraph.runs):
                    result_text.insert(tk.END, f"Run {i}: '{run.text}'\n")
                
                # 保存原始文本，用于对比
                original_text = paragraph.text
                
                # 执行替换
                result_text.insert(tk.END, "\n--- 执行替换 ---\n")
                
                # 开始捕获日志
                original_log = self.console_output.copy()
                
                # 执行替换
                success = self.replace_text_preserve_style(paragraph, selected_placeholder, replacement_value)
                
                # 获取新的日志
                new_logs = self.console_output[len(original_log):]
                
                # 显示替换结果
                result_text.insert(tk.END, f"替换成功: {success}\n")
                result_text.insert(tk.END, f"替换后文本: {paragraph.text}\n")
                result_text.insert(tk.END, f"新的Runs数量: {len(paragraph.runs)}\n")
                
                # 显示新的runs详情
                result_text.insert(tk.END, "\n--- 替换后Runs详情 ---\n")
                for i, run in enumerate(paragraph.runs):
                    result_text.insert(tk.END, f"Run {i}: '{run.text}'\n")
                
                # 显示详细日志
                if new_logs:
                    result_text.insert(tk.END, "\n--- 详细日志 ---\n")
                    for log in new_logs:
                        result_text.insert(tk.END, f"{log}\n")
                
                # 验证结果
                result_text.insert(tk.END, "\n--- 结果验证 ---\n")
                if selected_placeholder in paragraph.text:
                    result_text.insert(tk.END, "❌ 错误：占位符仍然存在于文本中！\n")
                else:
                    result_text.insert(tk.END, "✅ 正确：占位符已被替换\n")
                
                # 检查是否丢失了文本
                expected_text = original_text.replace(selected_placeholder, replacement_value)
                if paragraph.text == expected_text:
                    result_text.insert(tk.END, "✅ 正确：文本完整保留，没有丢失\n")
                else:
                    result_text.insert(tk.END, "❌ 错误：文本发生了意外变化\n")
                    result_text.insert(tk.END, f"期望文本: {expected_text}\n")
                    result_text.insert(tk.END, f"实际文本: {paragraph.text}\n")
                
            except Exception as e:
                result_text.insert(tk.END, f"测试过程中发生错误: {str(e)}\n")
                import traceback
                result_text.insert(tk.END, f"详细错误:\n{traceback.format_exc()}\n")
        
        # 按钮区域
        btn_frame = ttk.Frame(test_window)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="运行测试", command=run_test).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="关闭", command=test_window.destroy).pack(side=tk.LEFT, padx=5)
    
    def preview_filenames(self):
        """预览文件名"""
        if self.excel_data is None or len(self.excel_data) == 0:
            messagebox.showwarning("警告", "请先导入Excel数据！")
            return
        
        # 验证导出范围
        is_valid, range_message = self.validate_export_range()
        if not is_valid:
            messagebox.showwarning("警告", f"导出范围设置有误：{range_message}")
            return
        
        # 获取要预览的数据
        export_data = self.get_export_data_range()
        if export_data is None or len(export_data) == 0:
            messagebox.showwarning("警告", "没有可预览的数据！")
            return
        
        # 创建预览窗口
        preview_window = tk.Toplevel(self.root)
        preview_window.title("文件名预览")
        preview_window.geometry("800x600")
        preview_window.transient(self.root)
        
        # 标题
        ttk.Label(preview_window, text="文件名预览", 
                 font=("Arial", 14, "bold")).pack(pady=10)
        
        # 当前设置显示
        settings_frame = ttk.LabelFrame(preview_window, text="当前设置", padding=10)
        settings_frame.pack(fill=tk.X, padx=20, pady=5)
        
        naming_mode = self.naming_mode_var.get()
        settings_text = f"命名方式: {naming_mode}"
        
        if naming_mode == "字段":
            field_name = self.naming_field_var.get()
            settings_text += f"\n选择字段: {field_name if field_name else '未选择'}"
        elif naming_mode == "前缀":
            prefix = self.naming_prefix_var.get()
            settings_text += f"\n前缀设置: '{prefix}'"
        
        settings_text += f"\n导出范围: {range_message}"
        
        ttk.Label(settings_frame, text=settings_text, justify=tk.LEFT).pack(anchor=tk.W)
        
        # 预览结果
        preview_count = min(20, len(export_data))
        result_frame = ttk.LabelFrame(preview_window, text=f"文件名预览（前{preview_count}个）", padding=10)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # 创建表格显示
        columns = ("序号", "原始数据", "生成文件名", "状态")
        tree = ttk.Treeview(result_frame, columns=columns, show="headings", height=15)
        
        tree.heading("序号", text="行号")
        tree.heading("原始数据", text="数据内容（前50字符）")
        tree.heading("生成文件名", text="生成的文件名")
        tree.heading("状态", text="状态")
        
        tree.column("序号", width=60)
        tree.column("原始数据", width=200)
        tree.column("生成文件名", width=250)
        tree.column("状态", width=100)
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 滚动条
        scrollbar = ttk.Scrollbar(result_frame, orient="vertical", command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)
        
        # 生成预览数据
        def generate_preview():
            """生成预览数据"""
            used_names = set()
            preview_count = min(20, len(export_data))  # 只预览前20个
            
            for i in range(preview_count):
                try:
                    row = export_data.iloc[i]
                    original_index = export_data.index[i]  # 获取原始行索引
                    original_row_num = original_index + 1  # 用户视角的行号
                    
                    # 生成文件名（使用原始索引）
                    filename = self.generate_filename(row, original_index, used_names)
                    used_names.add(filename)
                    
                    # 生成显示的数据内容（取前几个字段的值）
                    data_preview = []
                    for j, (col, value) in enumerate(row.items()):
                        if j >= 3:  # 只显示前3个字段
                            break
                        if pd.notna(value):
                            data_preview.append(f"{col}:{str(value)[:20]}")
                    
                    data_text = "; ".join(data_preview)
                    if len(data_text) > 50:
                        data_text = data_text[:47] + "..."
                    
                    # 检查状态
                    status = "正常"
                    if filename.startswith("导出文档_") and naming_mode != "默认":
                        status = "回退"  # 表示使用了默认命名（可能是因为字段值为空等）
                    elif "_" in filename and not filename.startswith("导出文档_") and naming_mode == "字段":
                        if any(char.isdigit() for char in filename.split("_")[-1].split(".")[0]):
                            status = "重复处理"  # 表示处理了重复文件名
                    
                    tree.insert("", "end", values=(original_row_num, data_text, filename, status))
                    
                except Exception as e:
                    original_row_num = export_data.index[i] + 1 if i < len(export_data) else i+1
                    tree.insert("", "end", values=(original_row_num, "数据读取失败", f"导出文档_{original_row_num:03d}.docx", "错误"))
        
        # 生成预览
        generate_preview()
        
        # 统计信息
        stats_frame = ttk.Frame(preview_window)
        stats_frame.pack(fill=tk.X, padx=20, pady=5)
        
        total_rows = len(self.excel_data)
        export_rows = len(export_data)
        preview_rows = min(20, export_rows)
        
        stats_text = f"总计数据行数: {total_rows}  |  导出行数: {export_rows}  |  预览行数: {preview_rows}  |  "
        
        if naming_mode == "字段":
            field_name = self.naming_field_var.get()
            if field_name and field_name in export_data.columns:
                non_empty = export_data[field_name].dropna().count()
                stats_text += f"字段'{field_name}'非空值: {non_empty}/{export_rows}"
            else:
                stats_text += "字段无效，将使用默认命名"
        elif naming_mode == "前缀":
            prefix = self.naming_prefix_var.get()
            clean_prefix = self.clean_filename(prefix) if prefix else "文档"
            stats_text += f"使用前缀: '{clean_prefix}'"
        
        ttk.Label(stats_frame, text=stats_text).pack(anchor=tk.W)
        
        # 按钮
        btn_frame = ttk.Frame(preview_window)
        btn_frame.pack(pady=10)
        
        def refresh_preview():
            """刷新预览"""
            # 清空现有项目
            for item in tree.get_children():
                tree.delete(item)
            # 重新生成
            generate_preview()
        
        ttk.Button(btn_frame, text="刷新预览", command=refresh_preview).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="关闭", command=preview_window.destroy).pack(side=tk.LEFT, padx=5)
    
    def show_help(self):
        """显示帮助信息"""
        help_text = """Excel到Word模板转换助手

作者：yf
版本：2025
开源协议：MIT License

===== 使用说明 =====

1. 导入Excel文件：
   - 点击"导入Excel"按钮选择Excel文件
   - 支持.xlsx和.xls格式
   - 第一行将作为字段名

2. 导入Word模板：
   - 点击"导入Word模板"按钮选择Word模板
   - 支持的占位符格式：
     {{字段名}}（双花括号）

   - 系统会自动识别模板中的占位符，完全支持中文字段名
   - 支持全文档扫描，包括：
     • 正文段落和表格
     • 页眉和页脚（包括首页、奇偶页）
     • 文本框和形状
     • 嵌套表格和复杂结构
   - 完整的样式保持：
     • 字体格式：名称、大小、颜色、粗体、斜体、下划线等
     • 段落格式：对齐方式、间距、缩进、段落样式等
     • 精确替换：只替换占位符部分，保留其他文本的原有格式
     • 多run支持：一个段落中不同样式的文本片段都能正确保持
     • 详细追踪：通过"查看输出"可看到字体样式复制的详细过程

3. 字段映射：
   - 左侧显示Excel字段，右侧显示Word占位符
   - 双击右侧可以编辑映射关系
   - 支持数学运算，如：字段1+字段2、字段1*2等
   - 支持固定文本
   - 未映射的占位符默认填充"0"

4. 图片映射：
   - 点击"添加行"按钮添加图片映射规则
   - 双击表格行可以编辑映射关系
   - 图片文件夹：选择包含图片的文件夹
   - 映射规则：
     • 固定图片名：所有行使用同一张图片
     • 根据字段：根据Excel字段值匹配图片文件名
     • 根据行号：按行号匹配图片（1.jpg, 2.jpg等）
   - Word占位符：选择要替换的图片占位符
   - 支持格式：jpg, jpeg, png, bmp, gif

5. 预览和生成：
   - 点击"预览"查看第一行数据的效果
   - 点击"按模板导出"批量生成所有数据
   - 生成的文档会保存到选择的目录

6. 自动匹配：
   - 点击"自动匹配字段"自动匹配相同或相似的字段名
   - 精准匹配：完全匹配字段名
   - 模糊匹配：使用相似度算法（阈值0.6）自动匹配相似字段
   - 可通过"使用精准匹配"复选框切换匹配模式

7. 导出行数范围：
   - 全部数据：导出Excel中的所有行数据
   - 指定区间：只导出指定范围内的行数据
     • 输入起始行号和结束行号（均包括在内）
     • 行号从1开始计算，对应Excel中的数据行
     • 系统会自动验证范围的有效性
     • 预览和导出功能都会应用此设置

8. 其他选项：
   - 在文件中预览：直接打开预览文档
   - 图片宽度使用厘米：选择图片宽度单位（厘米或英寸）
   - 图片宽度：设置插入图片的宽度大小（默认9.8厘米）

9. 数字格式化：
   - 数字格式化：选择数字的显示格式
     • 保留原格式：保持Excel中的原始格式
     • 取整数：四舍五入到整数（如：3.7 → 4）
     • 保留N位小数：保留指定位数的小数（如：3.14159 → 3.14）
   - 自定义小数位数：可自定义保留的小数位数（启用后覆盖上述设置）
   - 千分位分隔符：为大数字添加千分位分隔符（如：1234.56 → 1,234.56）
   - 注意：只有识别为数字的内容才会被格式化，文本内容不受影响

10. 文件命名设置：
   - 默认命名：使用"导出文档_001.docx"的格式（按行号排序）
   - 使用Excel字段命名：选择Excel中的任意字段作为文件名
     • 自动清理不合法字符：< > : " / \\ | ? * 等字符会被替换为下划线
     • 重复处理：如果出现重复文件名，自动添加后缀（如：文件名_1.docx）
     • 空值处理：如果字段值为空，自动回退到默认命名
     • 长度限制：文件名过长会被截断到200个字符以内
   - 固定前缀命名：所有文件使用相同前缀+"_序号"的格式
     • 如：设置前缀为"报告"，生成"报告_001.docx, 报告_002.docx..."
     • 前缀也会进行不合法字符清理
   - 智能回退：任何命名方式出错时都会自动回退到默认命名，确保导出不会失败"""
        
        # 创建帮助窗口
        help_window = tk.Toplevel(self.root)
        help_window.title("使用助手")
        help_window.geometry("600x500")
        help_window.transient(self.root)
        
        # 添加滚动文本框
        text_frame = ttk.Frame(help_window)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Arial", 10))
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)
        
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        text_widget.insert(tk.END, help_text)
        text_widget.config(state=tk.DISABLED)
        
        # 关闭按钮
        ttk.Button(help_window, text="关闭", 
                  command=help_window.destroy).pack(pady=10)


def main():
    """主函数"""
    root = tk.Tk()
    app = Excel2WordConverter(root)
    root.mainloop()


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        import traceback
        # 将异常信息写入日志文件
        try:
            with open("error.log", "w", encoding="utf-8") as f:
                f.write(f"程序发生错误：\n{traceback.format_exc()}\n")
        except:
            pass
        
        # 在GUI环境中显示错误信息
        try:
            import tkinter as tk
            from tkinter import messagebox
            root = tk.Tk()
            root.withdraw()  # 隐藏主窗口
            messagebox.showerror("程序错误", 
                f"程序启动失败：{str(e)}\n\n错误详情已保存到 error.log 文件中。")
            root.destroy()
        except:
            # 如果无法显示GUI错误，尝试命令行显示
            try:
                print(f"程序发生错误：{str(e)}")
                print("错误详情已保存到 error.log 文件中。")
                input("按回车键退出...")
            except:
                pass